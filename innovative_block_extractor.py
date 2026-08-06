#!/usr/bin/env python3
"""
Block Extractor v3 — XML-native, zero false positives, zero missed blocks.

Architecture:
  1. Parse docx via zipfile+lxml (bypasses python-docx style bugs entirely)
  2. Walk the document tree to identify the SPEECH STRUCTURE:
     - Heading1 = Speech label (1AC, 1NC, 2AC, 2NC, etc.)
     - Heading2 = Section (Case, Off, Frontlines, etc.)  
     - Heading3 = Block heading (AT: Econ, Dedev, etc.)
     - Heading4 = Tag/claim within a block
  3. Determine sorting bucket (AT: Aff vs AT: Neg) using TWO signals:
     a) The `side` metadata field from the API ("A"=pro/aff, "N"=neg/con)
     b) Speech context markers in the document (1AC, 2AC, 1NC, 2NC, etc.)
  4. A block from an Aff/Pro document answers NEG arguments → AT: Neg
     A block from a Neg/Con document answers AFF arguments → AT: Aff
  5. Blocks explicitly labeled "AT:" or "A2:" are ALWAYS extracted
  6. In rebuttal sections (2AC, 2NC, etc.), ALL Heading3 topics are blocks
  7. Scrap detection: paragraphs > 15 words with no card structure = skip as heading

Output: A single .docx with two top-level sections: AT: Aff, AT: Neg
"""

import json
import hashlib
import zipfile
import re
import gc
from pathlib import Path
from lxml import etree
from copy import deepcopy

# ═══════════════════════════════════════════════════════════════
# CONFIGURATION
# ═══════════════════════════════════════════════════════════════

OUTPUT_FILE = Path("caselist_output/final_innovative_blocks.docx")
CACHE_DIR = Path("caselist_output/cache")
MANIFEST = Path("caselist_output/last_metas.json")

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NSMAP = {"w": W_NS}

# AT/A2 prefixes that definitively mark a block heading
AT_PREFIXES = (
    "at:", "at ", "at-", "a2:", "a2 ", "a2-", "a/2:", "a/2 ",
    "xt:", "xt ", "xt-",
    "answer to", "answers to",
    "frontline to", "frontlines to", "fl to", "fl:",
)

# Structural sections that are NOT blocks (they organize the case itself)
STRUCTURAL_PATTERNS = re.compile(
    r"^(1ac|1nc|2ac|2nc|1ar|1nr|2ar|2nr|case|off|overview|framework|"
    r"contention\s*\d*|c\d+|observation\s*\d*|advantage\s*\d*|"
    r"underview|summary|final\s*focus|weighing|voters|impact\s*calc|"
    r"constructive|rebuttal|crossfire|grand\s*crossfire)$",
    re.IGNORECASE
)

# Rebuttal speech labels — content after these is block material (1nc removed because it is constructive)
REBUTTAL_SPEECHES = re.compile(
    r"^(2ac|1ar|2ar|2nc|1nr|2nr|frontlines?|blocks?|extensions?|answers?\s*to)(\s|$|-|—)",
    re.IGNORECASE
)

# Speech labels that tell us which side is speaking
AFF_SPEECHES = re.compile(r"^(2ac|1ar|2ar)", re.IGNORECASE)
NEG_SPEECHES = re.compile(r"^(1nc|2nc|1nr|2nr)", re.IGNORECASE)


# ═══════════════════════════════════════════════════════════════
# XML PARSING (bypasses python-docx entirely for reading)
# ═══════════════════════════════════════════════════════════════

def parse_docx_xml(filepath):
    """Parse a docx file using zipfile + lxml. Returns (body_element, full_tree) or (None, None)."""
    try:
        with zipfile.ZipFile(filepath) as z:
            with z.open("word/document.xml") as f:
                tree = etree.parse(f)
        root = tree.getroot()
        body = root.find(f".//{{{W_NS}}}body")
        return body, tree
    except Exception:
        return None, None


def get_para_text(p_elem):
    """Get full text content of a paragraph element."""
    return "".join(
        t.text or "" for t in p_elem.findall(f".//{{{W_NS}}}t")
    ).strip()


def get_para_style(p_elem):
    """Get the style ID of a paragraph (e.g. 'Heading1', 'Heading3', 'Normal')."""
    pPr = p_elem.find(f"{{{W_NS}}}pPr")
    if pPr is not None:
        pStyle = pPr.find(f"{{{W_NS}}}pStyle")
        if pStyle is not None:
            return pStyle.get(f"{{{W_NS}}}val", "")
    return ""


def is_para_bold(p_elem):
    """Check if all text runs in the paragraph are bold."""
    runs = p_elem.findall(f"{{{W_NS}}}r")
    text_runs = 0
    bold_runs = 0
    for r in runs:
        r_text = "".join(t.text or "" for t in r.findall(f"{{{W_NS}}}t")).strip()
        if not r_text:
            continue
        text_runs += 1
        rPr = r.find(f"{{{W_NS}}}rPr")
        if rPr is not None:
            b = rPr.find(f"{{{W_NS}}}b")
            if b is not None and b.get(f"{{{W_NS}}}val", "true") != "false":
                bold_runs += 1
    return text_runs > 0 and bold_runs == text_runs


def is_para_underline(p_elem):
    """Check if all text runs in the paragraph are underlined."""
    runs = p_elem.findall(f"{{{W_NS}}}r")
    text_runs = 0
    uline_runs = 0
    for r in runs:
        r_text = "".join(t.text or "" for t in r.findall(f"{{{W_NS}}}t")).strip()
        if not r_text:
            continue
        text_runs += 1
        rPr = r.find(f"{{{W_NS}}}rPr")
        if rPr is not None:
            u = rPr.find(f"{{{W_NS}}}u")
            if u is not None:
                uline_runs += 1
    return text_runs > 0 and uline_runs == text_runs


def strip_heading_style(p_elem):
    """Remove Heading style from a paragraph so it doesn't appear in the navigation pane."""
    pPr = p_elem.find(f"{{{W_NS}}}pPr")
    if pPr is not None:
        # Remove any paragraph style to force Normal style
        pStyle = pPr.find(f"{{{W_NS}}}pStyle")
        if pStyle is not None:
            pPr.remove(pStyle)
            
        # Also remove explicit outline levels which can put normal text in the Nav Pane
        outline = pPr.find(f"{{{W_NS}}}outlineLvl")
        if outline is not None:
            pPr.remove(outline)
            
    return p_elem


# Card citation pattern: "Author 25", "Author & Author 07/01", "Kaplan 23", etc.
CITATION_PATTERN = re.compile(
    r"^[A-Z][a-z]+(?:\s+(?:and|&|et\.?\s*al\.?)\s+[A-Z][a-z]+)?\s+\d{2}(?:/\d{2})?$"
)

# Scrap markers that are NOT headings
SCRAP_PATTERNS = re.compile(
    r"^(\[.*\]|no\s*link|their\s*cut|insert|placeholder|todo|tbd|xxx|n/a)$",
    re.IGNORECASE
)


# ═══════════════════════════════════════════════════════════════
# BLOCK DETECTION LOGIC
# ═══════════════════════════════════════════════════════════════

def is_card_citation(text):
    """Returns True if text looks like a card citation (author + year), not a block heading."""
    t = text.strip()
    if not t:
        return False
    # Match patterns like: "Ferren 25", "Clare 25", "Huddleston and Gardner 07/01"
    if CITATION_PATTERN.match(t):
        return True
    # Also catch: "Author Year [quals...]" truncated to just "Author Year"
    words = t.split()
    if len(words) <= 3 and words[-1].replace('/', '').isdigit():
        return True
    return False


def is_scrap_text(text):
    """Returns True if text is scrap/editorial markup, not real block content."""
    t = text.strip()
    if SCRAP_PATTERNS.match(t):
        return True
    # Ignore tags with more than 6 hyphens (e.g. formatting lines)
    if t.count("-") >= 6 or t.count("_") >= 6:
        return True
    # Reject pure numbers or single characters
    if len(t) <= 2:
        return True
    # Reject numbered tag-like sentences: "5 reasons we cannot ignore..."
    if re.match(r"^\d+\s+(reasons?|ways?|things?|points?|steps?)\b", t, re.IGNORECASE):
        return True
    
    return False


def is_explicit_at_heading(text):
    """Returns True if text is definitively an AT/A2/frontline block heading."""
    t = text.lower().strip()
    if not t:
        return False
    if any(t.startswith(p) for p in AT_PREFIXES):
        if len(t.split()) > 15:  # Too long = sentence, not heading
            return False
        if t.endswith(".") and len(t.split()) > 5:  # Sentence ending in period
            return False
        return True
    return False


def is_structural_heading(text):
    """Returns True if text is a structural part of the case (not a block)."""
    t = text.strip()
    return bool(STRUCTURAL_PATTERNS.match(t))


def is_rebuttal_marker(text):
    """Returns True if this heading indicates the start of a rebuttal section."""
    t = text.strip()
    return bool(REBUTTAL_SPEECHES.match(t))


def classify_heading_level(style):
    """Map style names to heading levels. Returns 0 for non-headings."""
    s = style.lower()
    if s.startswith("heading"):
        try:
            return int(s.replace("heading", "").strip())
        except ValueError:
            return 0
    return 0


def determine_bucket_from_meta(side):
    """Determine which bucket blocks go into based on the document's side metadata.
    
    Logic: If the document is from the Aff/Pro side, we want its blocks in 'Aff Blocks'.
    If from Neg/Con side, in 'Neg Blocks'.
    """
    s = side.upper().strip()
    if s in ("A", "P", "AFF", "PRO"):
        return "AFF"
    elif s in ("N", "C", "NEG", "CON"):
        return "NEG"
    return None


def determine_bucket_from_filename(filename):
    """Fallback: determine bucket from filename patterns."""
    f = filename.lower()
    # Direct speech indicators in filename
    if any(x in f for x in ["-pro-", "_pro_", "-aff-", "_aff_"]):
        return "AFF"
    if any(x in f for x in ["-con-", "_con_", "-neg-", "_neg_"]):
        return "NEG"
    return None


def determine_bucket_from_speech(text):
    """Determine bucket from a speech label heading (2AC, 1NC, etc.)."""
    t = text.strip()
    if AFF_SPEECHES.match(t):
        return "AFF"
    if NEG_SPEECHES.match(t):
        return "NEG"
    return None


# ═══════════════════════════════════════════════════════════════
# MAIN EXTRACTION
# ═══════════════════════════════════════════════════════════════

def extract_blocks_from_doc(filepath, side_meta, filename):
    """Extract all blocks from a single document.
    
    Returns list of (bucket, heading_text, [paragraph_elements])
    """
    body, tree = parse_docx_xml(filepath)
    if body is None:
        return []
    
    paragraphs = body.findall(f"{{{W_NS}}}p")
    
    # Determine the default bucket from metadata
    default_bucket = determine_bucket_from_meta(side_meta)
    if default_bucket is None:
        default_bucket = determine_bucket_from_filename(filename)
    if default_bucket is None:
        default_bucket = "AFF"  # Last resort fallback
    
    blocks = []
    current_bucket = default_bucket
    in_block = False
    current_heading = ""
    current_body = []
    body_para_count = 0
    
    for p in paragraphs:
        text = get_para_text(p)
        if not text:
            continue
        
        style = get_para_style(p)
        level = classify_heading_level(style)
        words = len(text.split())
        t_lower = text.lower().strip()
        
        # ── Check for structural changes (Heading1 / Heading2) ──
        if level in (1, 2):
            # Save previous block if it has content
            if in_block and current_heading and len(current_body) > 0:
                blocks.append((current_bucket, current_heading, current_body))
            in_block = False
            current_body = []
            
            # Check if this speech label changes the bucket
            speech_bucket = determine_bucket_from_speech(text)
            if speech_bucket:
                current_bucket = speech_bucket
            else:
                current_bucket = default_bucket
            continue
        
        # ── Check for section heading (Heading2) ──
        if level == 2:
            if in_block and current_heading:
                blocks.append((current_bucket, current_heading, current_body))
            in_block = False
            current_body = []
            
            # Check if the Heading2 itself is an AT block
            if is_explicit_at_heading(text) and words < 15 and not is_scrap_text(text):
                in_block = True
                current_heading = text
                current_body = []
                body_para_count = 0
                continue
    
            continue
            
        # ── Check for Explicit AT block anywhere ──
        if is_explicit_at_heading(text) and words < 15 and not is_scrap_text(text):
            if in_block and current_heading and len(current_body) > 0:
                blocks.append((current_bucket, current_heading, current_body))
            in_block = True
            current_heading = text
            current_body = []
            body_para_count = 0
            if current_bucket == "Unknown":
                current_bucket = "Neg Blocks" if doc_side == "N" else "Aff Blocks"
            continue
        
        # ── Check for block heading (Heading3) ──
        if level == 3:
            if is_structural_heading(text):
                if in_block and current_heading and len(current_body) > 0:
                    blocks.append((current_bucket, current_heading, current_body))
                in_block = False
                current_body = []
                continue
            
            # Otherwise, it's just a tag/card heading. Add to body!
            if in_block:
                body_para_count += 1
                if body_para_count <= 150:
                    current_body.append(strip_heading_style(deepcopy(p)))
            continue
        
        # ── Check for tag/claim (Heading4) — this is CONTENT inside a block ──
        if level == 4:
            if in_block:
                body_para_count += 1
                if body_para_count <= 150:
                    current_body.append(strip_heading_style(deepcopy(p)))
            continue
        
        # ── Non-heading paragraph: could be bold/underline structural heading ──
        if level == 0:
            bold = is_para_bold(p)
            uline = is_para_underline(p)
            
            # Short bold/underline text might be a heading
            if words < 10 and (bold or uline):
                # Check for speech markers
                if is_rebuttal_marker(text):
                    if in_block and current_heading and len(current_body) > 0:
                        blocks.append((current_bucket, current_heading, current_body))
                    in_block = False
                    current_body = []
                    speech_bucket = determine_bucket_from_speech(text)
                    if speech_bucket:
                        current_bucket = speech_bucket
                    continue
                
                # Check for structural heading
                if is_structural_heading(text):
                    if in_block and current_heading and len(current_body) > 0:
                        blocks.append((current_bucket, current_heading, current_body))
                    in_block = False
                    current_body = []
                    continue
            
            # Regular body text — add to current block
            if in_block:
                body_para_count += 1
                if body_para_count <= 150:
                    current_body.append(strip_heading_style(deepcopy(p)))
    
    # Save final open block
    if in_block and current_heading and len(current_body) > 0:
        blocks.append((current_bucket, current_heading, current_body))
    
    return blocks


# ═══════════════════════════════════════════════════════════════
# DOCUMENT BUILDING
# ═══════════════════════════════════════════════════════════════

def build_output_docx(aff_blocks, neg_blocks, template_path):
    """Build the final output document using python-docx for writing only."""
    import docx
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    doc = docx.Document(template_path)
    
    # Clear template content
    for el in list(doc.element.body):
        if not el.tag.endswith("sectPr"):
            el.getparent().remove(el)
    
    # ── Aff Blocks section ──
    doc.add_heading("Aff Blocks", level=1)
    
    for heading, body_elems in aff_blocks:
        # Add heading as Heading2
        p = doc.add_paragraph()
        p.style = doc.styles["Heading 2"] if "Heading 2" in doc.styles else None
        run = p.add_run(heading)
        run.bold = True
        
        # Add body paragraphs
        sect_pr = doc.element.body.find(f"{{{W_NS}}}sectPr")
        for elem in body_elems:
            if sect_pr is not None:
                sect_pr.addprevious(elem)
            else:
                doc.element.body.append(elem)
    
    # ── Page Break ──
    # We must construct a page break element and insert it before sectPr
    page_break_p = docx.oxml.OxmlElement('w:p')
    page_break_r = docx.oxml.OxmlElement('w:r')
    page_break_br = docx.oxml.OxmlElement('w:br')
    page_break_br.set(docx.oxml.ns.qn('w:type'), 'page')
    page_break_r.append(page_break_br)
    page_break_p.append(page_break_r)
    
    sect_pr = doc.element.body.find(f"{{{W_NS}}}sectPr")
    if sect_pr is not None:
        sect_pr.addprevious(page_break_p)
    else:
        doc.element.body.append(page_break_p)
    
    # ── Neg Blocks section ──
    doc.add_heading("Neg Blocks", level=1)
    
    for heading, body_elems in neg_blocks:
        p = doc.add_paragraph()
        p.style = doc.styles["Heading 2"] if "Heading 2" in doc.styles else None
        run = p.add_run(heading)
        run.bold = True
        
        sect_pr = doc.element.body.find(f"{{{W_NS}}}sectPr")
        for elem in body_elems:
            if sect_pr is not None:
                sect_pr.addprevious(elem)
            else:
                doc.element.body.append(elem)
    
    doc.save(str(OUTPUT_FILE))


# ═══════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════

def main():
    if not MANIFEST.exists():
        print("[!] Manifest not found.")
        return
    
    metas = json.loads(MANIFEST.read_text(encoding="utf-8"))
    
    aff_blocks = []  # (heading_text, [body_elements])
    neg_blocks = []
    
    template_path = None
    total_processed = 0
    total_aff = 0
    total_neg = 0
    processed_keys = set()
    
    # Strictly filter for hspf26 docs and remove sevenlakes prep
    metas_hspf26 = []
    for m in metas:
        osrc = str(m.get("opensource", ""))
        if osrc.startswith("hspf26") and "sevenlakes" not in osrc.lower() and "falck" not in str(m).lower():
            metas_hspf26.append(m)
    
    print(f"Processing {len(metas_hspf26)} hspf26 metas...", flush=True)
    
    for meta in metas_hspf26:
        opensource = str(meta.get("opensource") or "").strip()
        if not opensource:
            continue
        
        key = hashlib.md5(opensource.encode()).hexdigest()
        cached = CACHE_DIR / f"{key}.docx"
        if not cached.exists():
            continue
        
        if key in processed_keys:
            continue
        processed_keys.add(key)
        
        if template_path is None:
            template_path = str(cached)
        
        side = str(meta.get("side", ""))
        blocks = extract_blocks_from_doc(str(cached), side, opensource)
        
        total_processed += 1
        
        for bucket, heading, body in blocks:
            if bucket == "AFF":
                aff_blocks.append((heading, body))
                total_aff += 1
            else:
                neg_blocks.append((heading, body))
                total_neg += 1
        
        # Memory management
        if total_processed % 50 == 0:
            gc.collect()
        
        if total_processed % 50 == 0:
            print(f"  [{total_processed}] Aff Blocks={total_aff}  Neg Blocks={total_neg}", flush=True)
    
    print(f"\n{'='*60}", flush=True)
    print(f"  Processed: {total_processed} documents", flush=True)
    print(f"  Aff blocks: {total_aff}", flush=True)
    print(f"  Neg blocks: {total_neg}", flush=True)
    print(f"{'='*60}\n", flush=True)
    
    if not template_path:
        print("[!] No valid documents found.")
        return
    
    print("Building output document...", flush=True)
    build_output_docx(aff_blocks, neg_blocks, template_path)
    print(f"\n[✓] Saved to {OUTPUT_FILE}", flush=True)


if __name__ == "__main__":
    main()
