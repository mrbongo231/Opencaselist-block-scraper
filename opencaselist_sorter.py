#!/usr/bin/env python3
"""
Standalone sorter for OpenCaselist manifests.

This script reads metadata from last_metas.json and DOCX bytes from cache,
then rebuilds a sorted packet using the bucket/speech/tournament algorithm.

Usage:
  python3 opencaselist_sorter.py
  python3 opencaselist_sorter.py --manifest caselist_output/last_metas.json --output caselist_output/compiled_blocks_sorted.docx
"""

from __future__ import annotations

import argparse
import copy
import hashlib
import io
import json
import os
import re
from collections import defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUTPUT_DIR = Path("caselist_output")
CACHE_DIR = OUTPUT_DIR / "cache"
DEFAULT_MANIFEST = OUTPUT_DIR / "last_metas.json"
DEFAULT_OUTPUT = OUTPUT_DIR / "compiled_blocks_sorted.docx"
DEFAULT_BLOCK_SPEECHES = ["2AC", "2NC", "1AR", "1NR", "2AR", "2NR", "Final Focus", "Crossfire/CX"]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Sort OpenCaselist blocks from manifest + cache")
    parser.add_argument("--manifest", type=Path, default=DEFAULT_MANIFEST, help="Path to metadata manifest JSON")
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT, help="Output sorted DOCX path")
    parser.add_argument(
        "--blocks-only",
        action="store_true",
        help="Keep only block-style speeches (2AC/2NC/1AR/1NR/2AR/2NR/Final Focus/Crossfire-CX)",
    )
    parser.add_argument(
        "--blocks-include-other",
        action="store_true",
        help="When --blocks-only is set, also keep sections labeled as 'Other'",
    )
    return parser.parse_args()


def classify_block_bucket(meta: dict) -> str:
    combined = " ".join(
        [
            str(meta.get("report") or ""),
            str(meta.get("opensource") or ""),
            str(meta.get("tournament") or ""),
        ]
    ).lower()

    compact = re.sub(r"[^a-z0-9]+", " ", combined)

    if re.search(r"\ba\s*2\s*aff\b", compact) or re.search(r"\ba2\s*aff\b", compact):
        return "A2 Aff"
    if re.search(r"\ba\s*2\s*neg\b", compact) or re.search(r"\ba2\s*neg\b", compact):
        return "A2 Neg"

    if re.search(r"\bpro\s*case\b", compact) or re.search(r"\baff\s*case\b", compact):
        return "Pro"
    if re.search(r"\bcon\s*case\b", compact) or re.search(r"\bneg\s*case\b", compact):
        return "Con"

    path = str(meta.get("opensource") or "").lower()
    if re.search(r"(^|[-_/])pro([-_/]|$)", path):
        return "Pro"
    if re.search(r"(^|[-_/])con([-_/]|$)", path):
        return "Con"

    side = str(meta.get("side") or "").upper()
    if side == "A":
        return "Pro"
    if side == "N":
        return "Con"

    return "Uncategorized"


def _add_bucket_heading(doc: Document, label: str):
    h = doc.add_heading(label, level=1)
    if h.runs:
        h.runs[0].font.color.rgb = RGBColor(0x11, 0x56, 0x99)


def _add_content_heading(doc: Document, label: str):
    h = doc.add_heading(label, level=2)
    if h.runs:
        h.runs[0].font.color.rgb = RGBColor(0x2A, 0x6A, 0xA8)


def _add_tournament_subheading(doc: Document, tourn_name: str):
    h = doc.add_heading(tourn_name, level=3)
    if h.runs:
        h.runs[0].font.color.rgb = RGBColor(0x1A, 0x5C, 0xA8)


SPEECH_ORDER = {
    "1AC": 10,
    "1NC": 20,
    "2AC": 30,
    "2NC": 40,
    "1AR": 50,
    "1NR": 60,
    "2AR": 70,
    "2NR": 80,
    "Final Focus": 90,
    "Crossfire/CX": 100,
    "Other": 999,
}


def extract_primary_speech_label(meta: dict) -> str:
    combined = " ".join(
        [
            str(meta.get("report") or ""),
            str(meta.get("opensource") or ""),
            str(meta.get("round") or ""),
        ]
    ).lower()
    compact = re.sub(r"[^a-z0-9]+", " ", combined)

    label_patterns = [
        ("1AC", [r"\b1ac\b", r"\bpc\b", r"\bpro\s*constructive\b"]),
        ("1NC", [r"\b1nc\b", r"\bcc\b", r"\bcon\s*constructive\b"]),
        ("2AC", [r"\b2ac\b", r"\bpr\b", r"\bpro\s*rebuttal\b"]),
        ("2NC", [r"\b2nc\b", r"\bcr\b", r"\bcon\s*rebuttal\b"]),
        ("1AR", [r"\b1ar\b"]),
        ("1NR", [r"\b1nr\b"]),
        ("2AR", [r"\b2ar\b", r"\bps\b", r"\bpro\s*summary\b"]),
        ("2NR", [r"\b2nr\b", r"\bcs\b", r"\bcon\s*summary\b"]),
        ("Final Focus", [r"\bfinal\s*focus\b", r"\bff\b", r"\bpf\b", r"\bcf\b"]),
        ("Crossfire/CX", [r"\bcross\s*(ex|examination|fire)\b", r"\bcx\b"]),
    ]

    for label, patterns in label_patterns:
        if any(re.search(p, compact) for p in patterns):
            return label

    return "Other"


def _round_sort_key(round_value: str):
    text = str(round_value or "").strip().lower()

    if text.isdigit():
        return (0, int(text), "")

    elim_order = {
        "doubles": 200,
        "triples": 210,
        "octas": 220,
        "quads": 230,
        "quarters": 240,
        "semis": 250,
        "finals": 260,
        "all": 999,
    }
    if text in elim_order:
        return (1, elim_order[text], "")

    num_match = re.search(r"\d+", text)
    if num_match:
        return (0, int(num_match.group()), text)

    return (2, 999, text)


def _demote_imported_heading_style(paragraph_element, paragraph_text=""):
    p_pr = paragraph_element.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        paragraph_element.insert(0, p_pr)

    p_style = p_pr.find(qn("w:pStyle"))
    if p_style is None:
        p_style = OxmlElement("w:pStyle")
        p_pr.insert(0, p_style)

    old_style_val = p_style.get(qn("w:val"), "")
    was_heading_style = re.fullmatch(r"Heading([1-9])", old_style_val, flags=re.IGNORECASE) is not None
    p_style.set(qn("w:val"), "Normal")

    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is not None:
        p_pr.remove(outline)

    if _looks_like_tag_or_card_heading(paragraph_text, was_heading_style):
        outline = OxmlElement("w:outlineLvl")
        outline.set(qn("w:val"), "6")
        p_pr.append(outline)


def _looks_like_tag_or_card_heading(text: str, was_heading_style: bool) -> bool:
    raw = (text or "").strip()
    if not raw or len(raw) > 180:
        return False

    compact = re.sub(r"[^a-z0-9: ]+", " ", raw.lower()).strip()
    if not compact:
        return False

    if re.search(r"\b(pro|con|aff|neg)\s*case\b", compact):
        return False
    if re.search(r"\baff\s*vs\b", compact):
        return False

    tag_patterns = [
        r"^\d+\s*[\]\).:-]\s*(at|a2|t|nl|nu|turn|ov|fw|link|impact|perm|alt|da|cp|k|solv|case|contention)\s*:",
        r"^(at|a2|t|nl|nu|turn|ov|fw|link|impact|perm|alt|da|cp|k|solv|case|contention)\s*:",
        r"^contention\s*\d+",
    ]
    if any(re.search(p, compact) for p in tag_patterns):
        return True

    if was_heading_style and ":" in raw and len(raw.split()) <= 12:
        return True

    return False


def _infer_side_code(meta: dict) -> str:
    path = str(meta.get("opensource") or "").lower()
    if re.search(r"(^|[-_/])pro([-_/]|$)", path):
        return "A"
    if re.search(r"(^|[-_/])con([-_/]|$)", path):
        return "N"

    side = str(meta.get("side") or "").upper()
    if side in ("A", "N"):
        return side

    compact = re.sub(
        r"[^a-z0-9]+",
        " ",
        " ".join([str(meta.get("report") or ""), str(meta.get("tournament") or "")]).lower(),
    )
    if re.search(r"\bpro\s*case\b|\baff\s*case\b", compact):
        return "A"
    if re.search(r"\bcon\s*case\b|\bneg\s*case\b", compact):
        return "N"
    return ""


def detect_speech_label_from_text(text: str, side_code: str) -> str:
    compact = re.sub(r"[^a-z0-9]+", " ", (text or "").lower()).strip()
    if not compact:
        return ""

    explicit = [
        ("1AC", [r"\b1ac\b", r"\bpc\b", r"\bpro\s*constructive\b"]),
        ("1NC", [r"\b1nc\b", r"\bcc\b", r"\bcon\s*constructive\b"]),
        ("2AC", [r"\b2ac\b", r"\bpr\b", r"\bpro\s*rebuttal\b"]),
        ("2NC", [r"\b2nc\b", r"\bcr\b", r"\bcon\s*rebuttal\b"]),
        ("1AR", [r"\b1ar\b"]),
        ("1NR", [r"\b1nr\b"]),
        ("2AR", [r"\b2ar\b", r"\bps\b", r"\bpro\s*summary\b"]),
        ("2NR", [r"\b2nr\b", r"\bcs\b", r"\bcon\s*summary\b"]),
        ("Final Focus", [r"\bfinal\s*focus\b", r"\bpf\b", r"\bcf\b", r"\bff\b"]),
        ("Crossfire/CX", [r"\bcross\s*(ex|examination|fire)\b", r"\bcx\b"]),
    ]
    for label, patterns in explicit:
        if any(re.search(p, compact) for p in patterns):
            return label

    if re.search(r"\b(rebuttal|block|frontline|a\s*2|a2)\b", compact):
        if side_code == "A":
            return "2AC"
        if side_code == "N":
            return "2NC"
        return "Other"

    if re.search(r"\b(constructive|case|pro\s*case|con\s*case)\b", compact):
        if side_code == "A":
            return "1AC"
        if side_code == "N":
            return "1NC"
        return "Other"

    if re.search(r"\bsummary\b", compact):
        if side_code == "A":
            return "2AR"
        if side_code == "N":
            return "2NR"
        return "Other"

    return ""


def _is_speech_boundary_heading(para) -> bool:
    text = (para.text or "").strip()
    if not text or len(text) > 120:
        return False

    try:
        style_name = (para.style.name or "").lower()
    except Exception:
        style_name = ""

    if style_name.startswith("heading"):
        return True

    if len(text.split()) <= 8 and re.fullmatch(r"[A-Za-z0-9\-: /&()]+", text):
        return True
    return False


def split_docx_into_speech_sections(src_bytes: bytes, meta: dict):
    try:
        src = Document(io.BytesIO(src_bytes))
    except Exception:
        return []

    side_code = _infer_side_code(meta)
    current_label = extract_primary_speech_label(meta)
    current_title = current_label
    current_paragraphs = []
    sections = []

    for para in src.paragraphs:
        text = (para.text or "").strip()
        detected = ""
        if text and _is_speech_boundary_heading(para):
            detected = detect_speech_label_from_text(text, side_code)

        if detected:
            if detected != current_label:
                if current_paragraphs:
                    sections.append((current_label, current_title, current_paragraphs))
                    current_paragraphs = []
                current_label = detected
                current_title = text[:120]
            elif current_title == current_label:
                current_title = text[:120]

        new_p = copy.deepcopy(para._element)
        _demote_imported_heading_style(new_p, text)
        current_paragraphs.append(new_p)

    if current_paragraphs:
        sections.append((current_label, current_title, current_paragraphs))

    return sections


def copy_section_into(paragraph_elements: list, dest_doc: Document) -> int:
    dest_body = dest_doc.element.body
    count = 0
    for el in paragraph_elements:
        new_el = copy.deepcopy(el)
        dest_body.append(new_el)
        count += 1

    dest_doc.add_paragraph()
    return count


def _section_entry_sort_key(meta: dict, section_index: int):
    return (_round_sort_key(str(meta.get("round") or "")), section_index, Path(str(meta.get("opensource") or "")).name.lower())


def normalize_speech_for_bucket(bucket: str, speech_label: str) -> str:
    if bucket in ("Con", "A2 Neg"):
        if speech_label == "1AC":
            return "1NC"
        if speech_label == "2AC":
            return "2NC"
    if bucket in ("Pro", "A2 Aff"):
        if speech_label == "1NC":
            return "1AC"
        if speech_label == "2NC":
            return "2AC"
    return speech_label


def build_allowed_block_speeches(include_other: bool) -> set[str]:
    allowed = set(DEFAULT_BLOCK_SPEECHES)
    if include_other:
        allowed.add("Other")
    return allowed


def build_cover(doc: Document, caselist: str, manifest_path: Path, loaded_count: int, non_docx: int, missing_cache: int):
    h0 = doc.add_heading("OpenCaselist Sorted Packet", 0)
    h0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if h0.runs:
        h0.runs[0].font.color.rgb = RGBColor(0x1A, 0x5F, 0xA8)

    doc.add_paragraph()

    rows = [
        ("Caselist", caselist),
        ("Source", str(manifest_path)),
        ("Loaded DOCX files", str(loaded_count)),
        ("Skipped non-DOCX", str(non_docx)),
        ("Missing cache DOCX", str(missing_cache)),
        ("Generated", datetime.now().strftime("%Y-%m-%d %H:%M")),
    ]

    for label, value in rows:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lb = p.add_run(f"{label}:  ")
        lb.bold = True
        lb.font.size = Pt(12)
        vv = p.add_run(value)
        vv.font.size = Pt(12)

    doc.add_page_break()


def _guess_caselist(metas: list[dict]) -> str:
    for meta in metas:
        path = str(meta.get("opensource") or "")
        parts = [p for p in path.split("/") if p]
        if parts:
            return parts[0]
    return os.getenv("CASELIST", "unknown")


def load_manifest(path: Path) -> list[dict]:
    try:
        data = json.loads(path.read_text())
    except Exception as e:
        raise SystemExit(f"Failed to read manifest: {e}")

    if not isinstance(data, list):
        raise SystemExit("Manifest must be a JSON list")
    return data


def load_cached_docx_bytes(opensource_path: str) -> bytes | None:
    key = hashlib.md5(opensource_path.encode()).hexdigest()
    cached = CACHE_DIR / f"{key}.docx"
    if not cached.exists():
        return None
    try:
        return cached.read_bytes()
    except Exception:
        return None


def main() -> None:
    args = parse_args()
    manifest_path = args.manifest
    output_path = args.output

    if not manifest_path.exists():
        raise SystemExit(f"Manifest not found: {manifest_path}")

    metas = load_manifest(manifest_path)
    if not metas:
        raise SystemExit("Manifest is empty")

    loaded = []
    non_docx = 0
    missing_cache = 0

    for meta in metas:
        path = str(meta.get("opensource") or "").strip()
        if not path:
            continue
        is_pdf = path.lower().endswith(".pdf")
        if not (path.lower().endswith(".docx") or is_pdf):
            non_docx += 1
            continue

        data = load_cached_docx_bytes(path)
        if data is None:
            missing_cache += 1
            continue

        loaded.append((meta, data))

    if not loaded:
        raise SystemExit("No cached DOCX files were loaded from manifest")

    out_doc = Document()
    for section in out_doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    caselist = _guess_caselist(metas)
    build_cover(out_doc, caselist, manifest_path, len(loaded), non_docx, missing_cache)

    by_bucket_speech_tourn = defaultdict(lambda: defaultdict(lambda: defaultdict(list)))

    allowed_block_speeches = build_allowed_block_speeches(args.blocks_include_other)
    if args.blocks_only:
        speech_order = ["1AC", "1NC", "2AC", "2NC", "1AR", "1NR", "2AR", "2NR", "Final Focus", "Crossfire/CX", "Other"]
        print("[->] Blocks-only mode enabled")
        print("[->] Keeping speeches:", ", ".join([s for s in speech_order if s in allowed_block_speeches]))

    for meta, data in loaded:
        bucket = classify_block_bucket(meta)
        tourn = str(meta.get("tournament") or "").lstrip("0123456789- ").strip() or "Unknown"
        sections = split_docx_into_speech_sections(data, meta)
        if not sections:
            sections = [(extract_primary_speech_label(meta), "Full File", [])]

        for idx, (speech_label, section_title, _) in enumerate(sections):
            speech_label = normalize_speech_for_bucket(bucket, speech_label)
            if args.blocks_only and speech_label not in allowed_block_speeches:
                continue
            by_bucket_speech_tourn[bucket][speech_label][tourn].append((meta, idx, section_title))

    bucket_order = ["A2 Aff", "A2 Neg", "Pro", "Con", "Uncategorized"]
    speech_order = ["1AC", "1NC", "2AC", "2NC", "1AR", "1NR", "2AR", "2NR", "Final Focus", "Crossfire/CX", "Other"]

    present_buckets = [b for b in bucket_order if b in by_bucket_speech_tourn]

    if not present_buckets:
        raise SystemExit("No sections available after filtering. Try without --blocks-only or with --blocks-include-other.")

    print("[->] Section counts:")
    for bucket in present_buckets:
        bucket_total = 0
        for tmap in by_bucket_speech_tourn[bucket].values():
            bucket_total += sum(len(v) for v in tmap.values())
        print(f"  - {bucket}: {bucket_total} sections")

    for bucket in present_buckets:
        _add_bucket_heading(out_doc, bucket)
        speech_map = by_bucket_speech_tourn[bucket]
        present_speeches = [s for s in speech_order if s in speech_map]

        for speech_label in present_speeches:
            _add_content_heading(out_doc, speech_label)
            tourn_map = speech_map[speech_label]

            for tourn_name in sorted(tourn_map.keys(), key=lambda x: str(x).lower()):
                entries = sorted(tourn_map[tourn_name], key=lambda row: _section_entry_sort_key(row[0], row[1]))
                _add_tournament_subheading(out_doc, tourn_name)
                for meta, sec_idx, section_title in entries:
                    # Parse the docx just-in-time to avoid holding 300+ parsed XML trees in memory
                    key = hashlib.md5(str(meta.get("opensource") or "").strip().encode()).hexdigest()
                    cached = Path("caselist_output/cache") / f"{key}.docx"
                    
                    try:
                        data = cached.read_bytes()
                        temp_sections = split_docx_into_speech_sections(data, meta)
                        if not temp_sections:
                            temp_sections = [(extract_primary_speech_label(meta), "Full File", [])]
                        
                        paragraph_elements = temp_sections[sec_idx][2] if sec_idx < len(temp_sections) else []
                        
                        n = copy_section_into(paragraph_elements, out_doc)
                        print(
                            f"  v  [{bucket} | {speech_label}] {Path(meta['opensource']).name} "
                            f"({n} paragraphs, section #{sec_idx + 1}: {section_title})"
                        )
                    except Exception as e:
                        print(f"  [!] Failed to re-parse {Path(meta.get('opensource', '')).name}: {e}")

        out_doc.add_page_break()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    out_doc.save(str(output_path))

    print("\n[v] Sorted DOCX saved:", output_path)
    print(f"[v] Loaded from cache: {len(loaded)}")
    print(f"[!] Skipped non-DOCX: {non_docx}")
    print(f"[!] Missing cache DOCX: {missing_cache}")


if __name__ == "__main__":
    main()
