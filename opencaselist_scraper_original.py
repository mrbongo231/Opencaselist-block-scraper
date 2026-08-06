#!/usr/bin/env python3
"""
Original-style OpenCaselist scraper (separate from sorter).

This script keeps the classic pipeline:
1) Resolve targets (teams/school/recent/topic)
2) Download files
3) Compile by tournament (no bucket/speech sorting)

The advanced sorter is provided in a separate script:
  opencaselist_sorter.py
"""

import copy
import hashlib
import io
import json
import os
import subprocess
import time
from collections import defaultdict
from datetime import datetime, timedelta, timezone
from pathlib import Path

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ============================
# Configuration
# ============================

CASELIST_TOKEN = (os.getenv("CASELIST_TOKEN", "").strip() or "7197e95921e0982ac01651ae3045ff26")
DEFAULT_CASELIST = "hspf26"
CASELIST = (os.getenv("CASELIST", DEFAULT_CASELIST).strip() or DEFAULT_CASELIST)

SPECIFIC_TEAMS = [
    ("StrakeJesuitCollegePreparatory", "CaMa"),
]

SPECIFIC_SCHOOLS = [
    "StrakeJesuitCollegePreparatory",
]

DAYS_RECENT = 7
TOPIC_KEYWORDS = []
RECENT_SINCE = os.getenv("RECENT_SINCE", "").strip()

OUTPUT_DIR = Path("caselist_output")
CACHE_DIR = OUTPUT_DIR / "cache"
OUTPUT_NAME = "compiled_blocks"
LAST_METAS_FILE = OUTPUT_DIR / "last_metas.json"
LAST_FAILED_FILE = OUTPUT_DIR / "last_failed_paths.json"

SCAN_LOG_EVERY_SCHOOLS = max(1, int(os.getenv("SCAN_LOG_EVERY_SCHOOLS", "10")))
SCAN_SHOW_TEAM = os.getenv("SCAN_SHOW_TEAM", "0").strip() == "1"
DOWNLOAD_LOG_EVERY = max(1, int(os.getenv("DOWNLOAD_LOG_EVERY", "25")))

OUTPUT_DIR.mkdir(exist_ok=True)
CACHE_DIR.mkdir(exist_ok=True)

API_BASE = "https://api.opencaselist.com/v1"

if not CASELIST_TOKEN or CASELIST_TOKEN == "REPLACE_ME_WITH_YOUR_TOKEN":
    raise SystemExit(
        "Missing CASELIST_TOKEN. Set it before running, e.g. "
        "export CASELIST_TOKEN='your_token_here'"
    )

session = requests.Session()
session.cookies.set("caselist_token", CASELIST_TOKEN, domain=".opencaselist.com")
session.headers.update(
    {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
        "Referer": "https://opencaselist.com/",
    }
)

AUTH_ERROR_REPORTED = False
TARGET_MODE = "teams"


# ============================
# Prompts
# ============================

def prompt_for_target_mode():
    print("\nSelect target mode:")
    print("  1. teams   - specific (school, team) pairs")
    print("  2. school  - all teams inside one or more schools")
    print("  3. recent  - rounds uploaded in last N days (site-wide scan)")
    print("  4. topic   - scan all teams matching topic keywords (site-wide scan)")
    print("  (Press Enter for default: teams)\n")

    choice = input("Enter choice (1-4): ").strip()

    if choice == "" or choice == "1":
        mode = "teams"
        teams = []
        print("\nEnter teams as: School,Team  (blank line to finish)")
        while True:
            line = input("Team: ").strip()
            if not line:
                break
            try:
                school, team = [x.strip() for x in line.split(",", 1)]
                if school and team:
                    teams.append((school, team))
                else:
                    print("  Must include both school and team.")
            except Exception:
                print("  Format must be: School,Team")

        if not teams:
            teams = SPECIFIC_TEAMS
        return mode, {"SPECIFIC_TEAMS": teams}

    if choice == "2":
        mode = "school"
        schools = []
        print("\nEnter school names (blank line to finish)")
        while True:
            line = input("School: ").strip()
            if not line:
                break
            schools.append(line)

        if not schools:
            schools = SPECIFIC_SCHOOLS
        return mode, {"SPECIFIC_SCHOOLS": schools}

    if choice == "3":
        mode = "recent"
        days = input("How many days back? (default 7): ").strip()
        try:
            days_int = int(days) if days else DAYS_RECENT
        except Exception:
            days_int = DAYS_RECENT
        return mode, {"DAYS_RECENT": days_int}

    if choice == "4":
        mode = "topic"
        print("\nEnter topic keywords separated by commas.")
        keywords = input("Keywords: ").strip()
        kw_list = [k.strip() for k in keywords.split(",") if k.strip()]
        return mode, {"TOPIC_KEYWORDS": kw_list}

    print("Invalid choice. Defaulting to teams.")
    return "teams", {"SPECIFIC_TEAMS": SPECIFIC_TEAMS}


def prompt_optional_topic_filter():
    print("\nOptional: apply a topic filter on top of your mode?")
    print("- Leave blank for NO filter (include all rounds).")
    print("- Or enter keywords separated by commas.")
    resp = input("Topic keywords (blank = none): ").strip()
    if not resp:
        return None
    return [k.strip() for k in resp.split(",") if k.strip()]


# ============================
# API Helpers
# ============================

def api_get(url, params=None, retries=3):
    global AUTH_ERROR_REPORTED

    for attempt in range(retries):
        try:
            r = session.get(url, params=params, timeout=15)
            if r.status_code in (401, 403):
                if not AUTH_ERROR_REPORTED:
                    print("  [auth] API returned 401/403 (Not Authorized).")
                    print("  [auth] CASELIST_TOKEN is invalid/expired or missing permissions.")
                    AUTH_ERROR_REPORTED = True
                return None
            if r.status_code == 429:
                wait = 10 * (attempt + 1)
                print(f"  [rate limit] waiting {wait}s...")
                time.sleep(wait)
                continue
            if r.status_code == 404:
                return None
            r.raise_for_status()
            return r.json()
        except Exception:
            if attempt == retries - 1:
                return None
            time.sleep(2 ** attempt)

    return None


def _extract_first_list(payload, preferred_keys=()):
    if isinstance(payload, list):
        return payload

    if not isinstance(payload, dict):
        return []

    for key in preferred_keys:
        val = payload.get(key)
        if isinstance(val, list):
            return val
        if isinstance(val, dict):
            for nested_key in ("items", "rows", "data", "results", "schools", "teams", "rounds"):
                nested_val = val.get(nested_key)
                if isinstance(nested_val, list):
                    return nested_val

    for val in payload.values():
        if isinstance(val, list):
            return val

    for val in payload.values():
        if isinstance(val, dict):
            for nested_val in val.values():
                if isinstance(nested_val, list):
                    return nested_val

    return []


def fetch_all_schools():
    print(f"[->] Fetching schools in {CASELIST}...")
    data = api_get(f"{API_BASE}/caselists/{CASELIST}/schools")
    if not data:
        return []
    schools = _extract_first_list(data, preferred_keys=("schools", "data", "items", "results"))
    print(f"    {len(schools)} schools found")
    return schools


def fetch_teams_in_school(school):
    data = api_get(f"{API_BASE}/caselists/{CASELIST}/schools/{school}/teams")
    if not data:
        return []
    return _extract_first_list(data, preferred_keys=("teams", "data", "items", "results"))


def fetch_rounds(school, team):
    cache_key = hashlib.md5(f"{CASELIST}{school}{team}".encode()).hexdigest()
    cache_file = CACHE_DIR / f"rounds_{cache_key}.json"
    if cache_file.exists() and (time.time() - cache_file.stat().st_mtime) < 3600:
        return json.loads(cache_file.read_text())

    data = api_get(f"{API_BASE}/caselists/{CASELIST}/schools/{school}/teams/{team}/rounds")
    if data is None:
        data = api_get(f"{API_BASE}/caselists/{CASELIST}/teams/{school}/{team}/rounds")
    if not data:
        return []

    rounds = _extract_first_list(data, preferred_keys=("rounds", "data", "items", "results"))
    cache_file.write_text(json.dumps(rounds))
    time.sleep(0.3)
    return rounds


# ============================
# Filters and Resolution
# ============================

def _parse_api_timestamp(raw_ts):
    text = (raw_ts or "").strip()
    if not text:
        return None
    normalized = text.replace("T", " ").replace("Z", "")
    normalized = normalized[:19]
    try:
        return datetime.strptime(normalized, "%Y-%m-%d %H:%M:%S")
    except Exception:
        return None


def _parse_since_timestamp(raw_ts):
    text = (raw_ts or "").strip()
    if not text:
        return None

    # Accept common forms for convenience.
    # Examples:
    #   2026-04-11 08:00
    #   2026-04-11 08:00:00
    #   2026-04-11T08:00
    #   2026-04-11T08:00:00
    candidates = [text, text.replace("T", " ")]
    formats = [
        "%Y-%m-%d %H:%M:%S",
        "%Y-%m-%d %H:%M",
        "%Y-%m-%d",
    ]

    for candidate in candidates:
        for fmt in formats:
            try:
                return datetime.strptime(candidate, fmt)
            except Exception:
                continue

    return None


def _is_recent(rnd, cutoff):
    dt = _parse_api_timestamp(rnd.get("created_at"))
    return dt is not None and dt >= cutoff


def _matches_topic(rnd):
    if not TOPIC_KEYWORDS:
        return True
    text = ((rnd.get("report") or "") + " " + (rnd.get("opensource") or "")).lower()
    return any(kw.lower() in text for kw in TOPIC_KEYWORDS)


def dedup_rounds(rounds):
    seen = {}
    for r in rounds:
        path = r.get("opensource")
        if path and path not in seen and _matches_topic(r):
            seen[path] = r
    return list(seen.values())


def resolve_targets():
    results = []

    if TARGET_MODE == "teams":
        total_teams = len(SPECIFIC_TEAMS)
        for idx, (school, team) in enumerate(SPECIFIC_TEAMS, start=1):
            print(f"[->] Team {idx}/{total_teams}: {school} / {team}")
            rounds = fetch_rounds(school, team)
            results.append((school, team, rounds))

    elif TARGET_MODE == "school":
        total_schools = len(SPECIFIC_SCHOOLS)
        scanned_teams = 0
        for school_idx, school in enumerate(SPECIFIC_SCHOOLS, start=1):
            print(f"[->] School {school_idx}/{total_schools}: {school}")
            teams = fetch_teams_in_school(school)
            print(f"    [scan] {len(teams)} teams found")
            for team_obj in teams:
                team = team_obj if isinstance(team_obj, str) else (team_obj.get("name") or team_obj.get("team", ""))
                if not team:
                    continue
                scanned_teams += 1
                rounds = fetch_rounds(school, team)
                results.append((school, team, rounds))
                time.sleep(0.2)
            print(
                f"    [progress] schools {school_idx}/{total_schools} | "
                f"teams scanned {scanned_teams} | teams resolved {len(results)}"
            )

    elif TARGET_MODE == "recent":
        if RECENT_SINCE:
            since_dt = _parse_since_timestamp(RECENT_SINCE)
            if since_dt is None:
                print(
                    "[!] Invalid RECENT_SINCE format. "
                    "Use YYYY-MM-DD HH:MM[:SS], YYYY-MM-DDTHH:MM[:SS], or YYYY-MM-DD."
                )
                return []
            cutoff = since_dt
            print(
                f"[->] Rounds uploaded since {cutoff.strftime('%Y-%m-%d %H:%M:%S')} "
                "(RECENT_SINCE override, created_at only)..."
            )
        else:
            cutoff = datetime.now(timezone.utc).replace(tzinfo=None) - timedelta(days=DAYS_RECENT)
            print(f"[->] Rounds uploaded since {cutoff.strftime('%Y-%m-%d')} ({DAYS_RECENT} days, created_at only)...")
        schools = fetch_all_schools()
        if not schools:
            print("[!] Could not fetch schools (token expired, rate-limited, or API temporarily unavailable).")
            return []
        total_schools = len(schools)
        scanned_schools = 0
        scanned_teams = 0
        matched_teams = 0
        for school_obj in schools:
            school = school_obj if isinstance(school_obj, str) else school_obj.get("name", "")
            if not school:
                continue
            scanned_schools += 1
            teams = fetch_teams_in_school(school)
            if scanned_schools % SCAN_LOG_EVERY_SCHOOLS == 0:
                print(
                    f"    [scan] school {scanned_schools}/{total_schools}: {school} "
                    f"({len(teams)} teams)",
                    flush=True,
                )

            for team_obj in teams:
                team = team_obj if isinstance(team_obj, str) else (team_obj.get("name") or team_obj.get("team", ""))
                if not team:
                    continue
                scanned_teams += 1
                if SCAN_SHOW_TEAM and scanned_teams % 25 == 0:
                    print(
                        f"      [team] {school} / {team} (team #{scanned_teams})",
                        flush=True,
                    )
                rounds = fetch_rounds(school, team)
                recent = [r for r in rounds if _is_recent(r, cutoff)]
                if recent:
                    results.append((school, team, recent))
                    matched_teams += 1

            if scanned_schools % 25 == 0 or scanned_schools == total_schools:
                print(
                    f"    [progress] schools {scanned_schools}/{total_schools} | "
                    f"teams scanned {scanned_teams} | matching teams {matched_teams}"
                )
            time.sleep(0.2)

        print(
            f"[->] Recent scan complete: {scanned_schools} schools, "
            f"{scanned_teams} teams scanned, {matched_teams} matching teams."
        )

    elif TARGET_MODE == "topic":
        if not TOPIC_KEYWORDS:
            print("[!] topic mode requires TOPIC_KEYWORDS to be set!")
            return []
        print(f"[->] Topic scan: {TOPIC_KEYWORDS}")
        schools = fetch_all_schools()
        if not schools:
            print("[!] Could not fetch schools (token expired, rate-limited, or API temporarily unavailable).")
            return []
        total_schools = len(schools)
        scanned_schools = 0
        scanned_teams = 0
        matched_teams = 0
        for school_obj in schools:
            school = school_obj if isinstance(school_obj, str) else school_obj.get("name", "")
            if not school:
                continue
            scanned_schools += 1
            teams = fetch_teams_in_school(school)
            if scanned_schools % SCAN_LOG_EVERY_SCHOOLS == 0:
                print(
                    f"    [scan] school {scanned_schools}/{total_schools}: {school} "
                    f"({len(teams)} teams)",
                    flush=True,
                )

            for team_obj in teams:
                team = team_obj if isinstance(team_obj, str) else (team_obj.get("name") or team_obj.get("team", ""))
                if not team:
                    continue
                scanned_teams += 1
                if SCAN_SHOW_TEAM and scanned_teams % 25 == 0:
                    print(
                        f"      [team] {school} / {team} (team #{scanned_teams})",
                        flush=True,
                    )
                rounds = fetch_rounds(school, team)
                matching = [r for r in rounds if _matches_topic(r)]
                if matching:
                    results.append((school, team, matching))
                    matched_teams += 1

            if scanned_schools % 25 == 0 or scanned_schools == total_schools:
                print(
                    f"    [progress] schools {scanned_schools}/{total_schools} | "
                    f"teams scanned {scanned_teams} | matches {matched_teams}"
                )
            time.sleep(0.2)

        print(
            f"[->] Topic scan complete: {scanned_schools} schools, "
            f"{scanned_teams} teams scanned, {matched_teams} matching teams."
        )

    return results


# ============================
# Download and Merge
# ============================

def download_file(path: str):
    if not path.lower().endswith(".docx"):
        print(f"    [!] Skipping non-DOCX source: {Path(path).name}")
        return None

    key = hashlib.md5(path.encode()).hexdigest()
    cached = CACHE_DIR / f"{key}.docx"
    if cached.exists():
        return cached.read_bytes()

    print(f"    [v] {Path(path).name}")
    for attempt in range(3):
        try:
            r = session.get(f"{API_BASE}/download", params={"path": path}, timeout=30)
            if r.status_code == 200 and r.content[:4] == b"PK\x03\x04":
                cached.write_bytes(r.content)
                time.sleep(0.6)
                return r.content
            time.sleep(2 ** attempt)
        except Exception:
            time.sleep(2 ** attempt)

    print(f"    [!] Failed after 3 attempts: {Path(path).name}")
    return None


def _add_attr_paragraph(doc, text, hex_color, bold=False, size_pt=10, space_before_pt=0, space_after_pt=3):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before_pt)
    para.paragraph_format.space_after = Pt(space_after_pt)
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))
    return para


def _add_rule(doc, color="3366AA"):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    p_pr = para._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    p_bdr.append(bot)
    p_pr.append(p_bdr)


def copy_docx_into(src_bytes: bytes, dest_doc: Document, meta: dict) -> int:
    try:
        src = Document(io.BytesIO(src_bytes))
    except Exception as e:
        print(f"    [!] Parse error: {e}")
        return 0

    side = "AFF" if meta.get("side") == "A" else "NEG"
    tourn = str(meta.get("tournament", "")).lstrip("0123456789- ").strip()
    rnd = str(meta.get("round", ""))
    opp = str(meta.get("opponent", ""))
    judge = str(meta.get("judge", ""))
    fname = Path(str(meta.get("opensource", ""))).name

    _add_attr_paragraph(dest_doc, "-" * 72, "2255AA", size_pt=7, space_before_pt=12, space_after_pt=1)
    _add_attr_paragraph(
        dest_doc,
        f"{meta.get('school', '')} / {meta.get('team', '')}   {side}   {tourn}  Round {rnd}",
        "1a5fa8",
        bold=True,
        size_pt=11,
        space_before_pt=1,
        space_after_pt=1,
    )
    if opp:
        _add_attr_paragraph(dest_doc, f"vs {opp}   |   Judge: {judge}", "777777", size_pt=9, space_after_pt=1)
    report = str(meta.get("report", ""))
    if report:
        _add_attr_paragraph(dest_doc, report.replace("\n", "  |  "), "999999", size_pt=8, space_after_pt=1)
    _add_attr_paragraph(dest_doc, f"File: {fname}", "AAAAAA", size_pt=7, space_after_pt=3)
    _add_rule(dest_doc)

    dest_body = dest_doc.element.body
    insert_idx = len(dest_body) - 1
    count = 0
    for para in src.paragraphs:
        new_p = copy.deepcopy(para._element)
        dest_body.insert(insert_idx, new_p)
        insert_idx += 1
        count += 1

    dest_doc.add_paragraph()
    return count


def build_cover(doc, target_summary, file_count, topic_info):
    h0 = doc.add_heading("OpenCaselist Block Compilation", 0)
    h0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if h0.runs:
        h0.runs[0].font.color.rgb = RGBColor(0x1A, 0x5F, 0xA8)

    doc.add_paragraph()

    for label, value in [
        ("Caselist", CASELIST),
        ("Mode", TARGET_MODE),
        ("Targets", target_summary),
        ("Files", f"{file_count} unique round documents"),
        ("Topic filter", topic_info),
        ("Generated", datetime.now().strftime("%Y-%m-%d %H:%M")),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lb = p.add_run(f"{label}:  ")
        lb.bold = True
        lb.font.size = Pt(12)
        val = p.add_run(str(value))
        val.font.size = Pt(12)

    doc.add_page_break()


def convert_to_pdf(docx_path: Path):
    pdf_path = docx_path.with_suffix(".pdf")

    try:
        from docx2pdf import convert

        print("[->] Converting to PDF via Microsoft Word (docx2pdf)...")
        convert(str(docx_path), str(pdf_path))
        if pdf_path.exists():
            print(f"[v] PDF saved: {pdf_path.resolve()}")
            return
    except Exception as e:
        print(f"[!] docx2pdf error: {e}")

    for cmd in ["soffice", "libreoffice", r"C:\\Program Files\\LibreOffice\\program\\soffice.exe"]:
        try:
            res = subprocess.run(
                [cmd, "--headless", "--convert-to", "pdf", "--outdir", str(docx_path.parent), str(docx_path)],
                capture_output=True,
                timeout=120,
            )
            if res.returncode == 0 and pdf_path.exists():
                print(f"[v] PDF saved via LibreOffice: {pdf_path.resolve()}")
                return
        except Exception:
            continue

    print("\n" + "=" * 55)
    print("  DOCX saved but PDF conversion unavailable.")
    print("  To get a PDF, either:")
    print("    1. pip install docx2pdf  (needs MS Word)")
    print("    2. Open the .docx in Word and Save As PDF")
    print(f"\n  DOCX is at: {docx_path.resolve()}")
    print("=" * 55 + "\n")


# ============================
# Main
# ============================

def main():
    global TARGET_MODE, SPECIFIC_TEAMS, SPECIFIC_SCHOOLS, DAYS_RECENT, TOPIC_KEYWORDS

    print("\n" + "=" * 60)
    print("  OpenCaselist Scraper v2 (Original Flow)")
    print(f"  caselist={CASELIST}")
    print("=" * 60)

    TARGET_MODE, updates = prompt_for_target_mode()
    for key, value in updates.items():
        globals()[key] = value

    extra_filter = prompt_optional_topic_filter()
    if extra_filter is not None:
        TOPIC_KEYWORDS = extra_filter

    print(f"\n[->] Running with mode={TARGET_MODE}")
    if TOPIC_KEYWORDS:
        print(f"[->] Topic filter: {TOPIC_KEYWORDS}")

    team_data = resolve_targets()
    if not team_data:
        print("[!] No targets resolved. Check configuration/mode.")
        return
    print(f"\n[v] {len(team_data)} teams resolved\n")

    all_metas = []
    for school, team, rounds in team_data:
        unique = dedup_rounds(rounds)
        print(f"  {school}/{team}: {len(unique)} unique files")
        for rnd in unique:
            path = rnd.get("opensource")
            if not path:
                continue
            all_metas.append(
                {
                    "school": school,
                    "team": team,
                    "tournament": rnd.get("tournament", ""),
                    "round": rnd.get("round", ""),
                    "side": rnd.get("side", ""),
                    "opponent": rnd.get("opponent", ""),
                    "judge": rnd.get("judge", ""),
                    "report": rnd.get("report", ""),
                    "opensource": path,
                    "created_at": rnd.get("created_at", ""),
                }
            )

    print(f"\n[->] Downloading {len(all_metas)} files...\n")
    LAST_METAS_FILE.write_text(json.dumps(all_metas, indent=2))
    print(f"[->] Saved manifest: {LAST_METAS_FILE.resolve()}")

    downloaded = []
    failed = []
    total_to_download = len(all_metas)
    for idx, meta in enumerate(all_metas, start=1):
        data = download_file(meta["opensource"])
        if data:
            downloaded.append((meta, data))
        else:
            failed.append(meta["opensource"])
        if idx % DOWNLOAD_LOG_EVERY == 0 or idx == total_to_download:
            print(
                f"    [download] {idx}/{total_to_download} | "
                f"ok {len(downloaded)} | failed {len(failed)}",
                flush=True,
            )

    LAST_FAILED_FILE.write_text(json.dumps(sorted(set(failed)), indent=2))
    print(f"[->] Failed list saved: {LAST_FAILED_FILE.resolve()}")

    print(f"\n[v] {len(downloaded)} files ready\n")
    if not downloaded:
        print("[!] Nothing to compile.")
        return

    out_doc = Document()
    for section in out_doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    target_summary = ", ".join(f"{s}/{t}" for s, t, _ in team_data) if len(team_data) <= 5 else f"{len(team_data)} teams"
    topic_info = " | ".join(TOPIC_KEYWORDS) if TOPIC_KEYWORDS else "none (all rounds included)"
    build_cover(out_doc, target_summary, len(downloaded), topic_info)

    by_tourn = defaultdict(list)
    for meta, data in downloaded:
        tourn = str(meta.get("tournament", "")).lstrip("0123456789- ").strip() or "Unknown"
        by_tourn[tourn].append((meta, data))

    for tourn_name in sorted(by_tourn.keys(), key=lambda x: str(x).lower()):
        h = out_doc.add_heading(tourn_name, level=1)
        if h.runs:
            h.runs[0].font.color.rgb = RGBColor(0x1A, 0x5C, 0xA8)

        for meta, data in by_tourn[tourn_name]:
            n = copy_docx_into(data, out_doc, meta)
            print(f"  v  {Path(meta['opensource']).name}  ({n} paragraphs)")

        out_doc.add_page_break()

    docx_path = OUTPUT_DIR / f"{OUTPUT_NAME}.docx"
    out_doc.save(str(docx_path))
    print(f"\n[v] DOCX saved: {docx_path.resolve()}")

    convert_to_pdf(docx_path)

    print("\n" + "=" * 60)
    print(f"  Done! Folder: {OUTPUT_DIR.resolve()}")
    print("=" * 60 + "\n")


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n[!] Cancelled by user (Ctrl+C). Exiting cleanly.")
