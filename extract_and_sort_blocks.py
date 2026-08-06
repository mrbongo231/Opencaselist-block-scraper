#!/usr/bin/env python3
import re
import copy
import sys
import os
from docx import Document

import glob

OUTPUT_FILE = "caselist_output/compiled_blocks_only_fixed.docx"

# Find all part files if the main file didn't merge successfully
source_files = sorted(glob.glob("caselist_output/compiled_smart_sorted_part*.docx"))
if os.path.exists("caselist_output/compiled_smart_sorted.docx"):
    source_files = ["caselist_output/compiled_smart_sorted.docx"]

if not source_files:
    print("Error: No compiled_smart_sorted*.docx files found.")
    sys.exit(1)

print(f"Loading {len(source_files)} compiled document(s)...")

def create_blank_doc(template_path):
    doc = Document(template_path)
    for el in doc.element.body:
        if el.tag.endswith('sectPr'):
            continue
        el.getparent().remove(el)
    return doc

print("Initializing output documents...")
doc_aff = create_blank_doc(source_files[0])
doc_neg = create_blank_doc(source_files[0])

AT_PATTERNS = [
    re.compile(r"^\s*AT\s*[:\-]?\s+", re.IGNORECASE),
    re.compile(r"^\s*A\s*[/-]?\s*2\s*[:\-]?\s+", re.IGNORECASE)
]

def is_at_block(text):
    for pat in AT_PATTERNS:
        if pat.match(text):
            return True
    return False

in_block = False
current_h1_text = ""
current_h3 = None
current_h4 = None

# Track whether we've emitted the current h3/h4 to the target doc to avoid duplicates
emitted_h3_aff = None
emitted_h4_aff = None
emitted_h3_neg = None
emitted_h4_neg = None

added_aff = 0
added_neg = 0

print("Extracting and sorting blocks...")
for sf in source_files:
    print(f"Processing {sf}...")
    src_doc = Document(sf)
    for p in src_doc.paragraphs:
        style = p.style.name if p.style else ""
        text = (p.text or "").strip()
        
        if style == "Heading 1":
            current_h1_text = text
            in_block = False
        elif style == "Heading 2":
            in_block = False
        elif style == "Heading 3":
            current_h3 = p
            current_h4 = None
            in_block = False
        elif style == "Heading 4":
            current_h4 = p
            in_block = False
        elif style == "Heading 5":
            if is_at_block(text):
                in_block = True
                
                # Pro Case / AT: Neg -> AT: Neg
                # Con Case / AT: Aff -> AT: Aff
                if "Pro" in current_h1_text or "Neg" in current_h1_text:
                    target_doc = doc_neg
                    is_aff = False
                else:
                    target_doc = doc_aff
                    is_aff = True
                    
                # Emit file headings if we haven't already for this side/file combination
                if is_aff:
                    if current_h3 and emitted_h3_aff != current_h3.text:
                        target_doc.element.body.append(copy.deepcopy(current_h3._element))
                        emitted_h3_aff = current_h3.text
                    if current_h4 and emitted_h4_aff != current_h4.text:
                        target_doc.element.body.append(copy.deepcopy(current_h4._element))
                        emitted_h4_aff = current_h4.text
                else:
                    if current_h3 and emitted_h3_neg != current_h3.text:
                        target_doc.element.body.append(copy.deepcopy(current_h3._element))
                        emitted_h3_neg = current_h3.text
                    if current_h4 and emitted_h4_neg != current_h4.text:
                        target_doc.element.body.append(copy.deepcopy(current_h4._element))
                        emitted_h4_neg = current_h4.text
                    
                target_doc.element.body.append(copy.deepcopy(p._element))
                if is_aff: added_aff += 1
                else: added_neg += 1
            else:
                in_block = False
        else:
            # Normal paragraphs: append to the target doc if we're currently inside a block
            if in_block:
                if "Pro" in current_h1_text or "Neg" in current_h1_text:
                    doc_neg.element.body.append(copy.deepcopy(p._element))
                    added_neg += 1
                else:
                    doc_aff.element.body.append(copy.deepcopy(p._element))
                    added_aff += 1

print(f"Extracted {added_aff} elements for AT: AFF")
print(f"Extracted {added_neg} elements for AT: NEG")

# Combine them into a final document
final_doc = create_blank_doc(source_files[0])

print("Combining into final document...")

# Add AT: AFF Heading 1
h1_aff = final_doc.add_heading('AT: Aff', level=1)
h1_aff.style = 'Heading 1'

for el in doc_aff.element.body:
    if not el.tag.endswith('sectPr'):
        final_doc.element.body.append(copy.deepcopy(el))

final_doc.add_page_break()

# Add AT: NEG Heading 1
h1_neg = final_doc.add_heading('AT: Neg', level=1)
h1_neg.style = 'Heading 1'

for el in doc_neg.element.body:
    if not el.tag.endswith('sectPr'):
        final_doc.element.body.append(copy.deepcopy(el))

final_doc.save(OUTPUT_FILE)
print(f"Saved to {OUTPUT_FILE}")
