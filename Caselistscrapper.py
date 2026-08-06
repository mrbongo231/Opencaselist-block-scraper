"""
OpenCaselist Scraper v2 (Interactive Target Mode)
=================================================
Downloads open-source round files and compiles them into a PDF,
preserving the ORIGINAL formatting from the source documents.

QUICK START:
  1. pip install requests python-docx docx2pdf
    2. Set env var CASELIST_TOKEN (DO NOT SHARE IT)
  3. Run: python caselist_scraper.py
  4. Choose a TARGET_MODE when prompted

TARGET_MODE options:
  "teams"   - specific list of (school, team) pairs you name
  "school"  - every team inside one or more schools
  "recent"  - rounds uploaded in the last N days (across whole caselist)
  "topic"   - scan all teams, include only rounds matching topic keywords

TOPIC FILTERING:
  Set TOPIC_KEYWORDS to words/phrases that appear in round reports.
  A round is included if its report contains ANY keyword.
  Set to [] to include everything.

PDF CONVERSION:
  On Windows with Microsoft Word: uses docx2pdf (best quality).
  Without Word: install LibreOffice and it auto-detects it,
  OR open the saved .docx in Word and Save As PDF manually.
"""

import requests
import hashlib
import time
import io
import os
import copy
import json
import subprocess
import re
from pathlib import Path
from datetime import datetime, timedelta, timezone
from collections import defaultdict

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ═══════════════════════════════════════════════════════════════
#  CONFIGURATION — EDIT THIS SECTION
# ═══════════════════════════════════════════════════════════════

# Read token from environment variable so secrets are not hardcoded.
CASELIST_TOKEN = (os.getenv("CASELIST_TOKEN", "").strip() or "7197e95921e0982ac01651ae3045ff26")
DEFAULT_CASELIST = "hspf26"
CASELIST = (os.getenv("CASELIST", DEFAULT_CASELIST).strip() or DEFAULT_CASELIST)

# Defaults (used if you choose to keep them / fallbacks)
SPECIFIC_TEAMS = [
    ("StrakeJesuitCollegePreparatory", "CaMa"),
    # ("Lexington", "MS"),
    # ("Westwood", "AG"),
]

SPECIFIC_SCHOOLS = [
    "StrakeJesuitCollegePreparatory",
    # "Lexington",
]

DAYS_RECENT = 7

# Topic keyword filter — applies on top of any mode above.
TOPIC_KEYWORDS = []

# Output settings
OUTPUT_DIR  = Path("caselist_output")
CACHE_DIR   = OUTPUT_DIR / "cache"
OUTPUT_NAME = "compiled_blocks"
LAST_METAS_FILE = OUTPUT_DIR / "last_metas.json"
LAST_FAILED_FILE = OUTPUT_DIR / "last_failed_paths.json"

# Scan logging controls (set via environment variables)
SCAN_LOG_EVERY_SCHOOLS = max(1, int(os.getenv("SCAN_LOG_EVERY_SCHOOLS", "10")))
SCAN_SHOW_TEAM = os.getenv("SCAN_SHOW_TEAM", "0").strip() == "1"
RECENT_MAX_PAGES = max(1, int(os.getenv("RECENT_MAX_PAGES", "20")))
DOWNLOAD_LOG_EVERY = max(1, int(os.getenv("DOWNLOAD_LOG_EVERY", "25")))
DOWNLOAD_ONLY = os.getenv("DOWNLOAD_ONLY", "0").strip() == "1"
DOWNLOAD_FAILED_ONLY = os.getenv("DOWNLOAD_FAILED_ONLY", "0").strip() == "1"

# ═══════════════════════════════════════════════════════════════

OUTPUT_DIR.mkdir(exist_ok=True)
CACHE_DIR.mkdir(exist_ok=True)

API_BASE = "https://api.opencaselist.com/v1"

if not CASELIST_TOKEN or CASELIST_TOKEN == "REPLACE_ME_WITH_YOUR_TOKEN":
    raise SystemExit(
        "Missing CASELIST_TOKEN. Set it before running, e.g. "
        "export CASELIST_TOKEN='your_token_here'"
    )

session = requests.Session()
session.cookies.set("caselist_token", CASELIST_TOKEN, domain=".opencaselist.com")
session.headers.update({
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
    "Referer": "https://opencaselist.com/",
})

AUTH_ERROR_REPORTED = False


# ───────────────────────────────────────────────────────────────
#  INTERACTIVE TARGET MODE PROMPT
# ───────────────────────────────────────────────────────────────

def prompt_for_target_mode():
    print("\nSelect target mode:")
    print("  1. teams   - specific (school, team) pairs")
    print("  2. school  - all teams inside one or more schools")
    print("  3. recent  - rounds uploaded in last N days (site-wide scan)")
    print("  4. topic   - scan all teams matching topic keywords (site-wide scan)")
    print("  (Press Enter for default: teams)\n")

    choice = input("Enter choice (1-4): ").strip()

    if choice == "" or choice == "1":
        mode = "teams"
        teams = []
        print("\nEnter teams as: School,Team  (blank line to finish)")
        print("Example: StrakeJesuitCollegePreparatory,CaMa\n")
        while True:
            line = input("Team: ").strip()
            if not line:
                break
            try:
                school, team = [x.strip() for x in line.split(",", 1)]
                if school and team:
                    teams.append((school, team))
                else:
                    print("  Must include both school and team.")
            except Exception:
                print("  Format must be: School,Team")

        # fallback to config defaults if user entered nothing
        if not teams:
            teams = SPECIFIC_TEAMS

        return mode, {"SPECIFIC_TEAMS": teams}

    if choice == "2":
        mode = "school"
        schools = []
        print("\nEnter school names (blank line to finish)")
        print("Example: StrakeJesuitCollegePreparatory\n")
        while True:
            line = input("School: ").strip()
            if not line:
                break
            schools.append(line)

        if not schools:
            schools = SPECIFIC_SCHOOLS

        return mode, {"SPECIFIC_SCHOOLS": schools}

    if choice == "3":
        mode = "recent"
        days = input("How many days back? (default 7): ").strip()
        try:
            days_int = int(days) if days else DAYS_RECENT
        except Exception:
            days_int = DAYS_RECENT
        return mode, {"DAYS_RECENT": days_int}

    if choice == "4":
        mode = "topic"
        print("\nEnter topic keywords separated by commas.")
        print('Example: antitrust, FTC, data, AI, tech, court\n')
        keywords = input("Keywords: ").strip()
        kw_list = [k.strip() for k in keywords.split(",") if k.strip()]
        return mode, {"TOPIC_KEYWORDS": kw_list}

    print("Invalid choice. Defaulting to teams.")
    return "teams", {"SPECIFIC_TEAMS": SPECIFIC_TEAMS}


def prompt_optional_topic_filter():
    """
    Optional topic filter that can apply on top of any mode (including teams/school/recent).
    If the user leaves blank, keep current TOPIC_KEYWORDS (usually []).
    """
    print("\nOptional: apply a topic filter on top of your mode?")
    print("- Leave blank for NO filter (include all rounds).")
    print("- Or enter keywords separated by commas.")
    print('  Example: K, framework, ontology\n')
    resp = input("Topic keywords (blank = none): ").strip()
    if not resp:
        return None
    kw_list = [k.strip() for k in resp.split(",") if k.strip()]
    return kw_list


# ───────────────────────────────────────────────────────────────
#  API HELPERS
# ───────────────────────────────────────────────────────────────

def api_get(url, params=None, retries=3):
    global AUTH_ERROR_REPORTED
    for attempt in range(retries):
        try:
            r = session.get(url, params=params, timeout=15)
            if r.status_code in (401, 403):
                if not AUTH_ERROR_REPORTED:
                    print("  [auth] API returned 401/403 (Not Authorized).")
                    print("  [auth] CASELIST_TOKEN in this shell is invalid/expired or missing permissions.")
                    print("  [auth] Re-export CASELIST_TOKEN, then rerun the scraper.")
                    AUTH_ERROR_REPORTED = True
                return None
            if r.status_code == 429:
                wait = 10 * (attempt + 1)
                print(f"  [rate limit] waiting {wait}s...")
                time.sleep(wait)
                continue
            if r.status_code == 404:
                return None
            r.raise_for_status()
            return r.json()
        except Exception:
            if attempt == retries - 1:
                return None
            time.sleep(2 ** attempt)
    return None


def _extract_first_list(payload, preferred_keys=()):
    """
    Extract the first list-like collection from heterogeneous API payloads.
    This keeps the scraper resilient if endpoint wrappers change.
    """
    if isinstance(payload, list):
        return payload

    if not isinstance(payload, dict):
        return []

    # First, honor caller-preferred keys.
    for key in preferred_keys:
        val = payload.get(key)
        if isinstance(val, list):
            return val
        if isinstance(val, dict):
            for nested_key in ("items", "rows", "data", "results", "schools", "teams", "rounds"):
                nested_val = val.get(nested_key)
                if isinstance(nested_val, list):
                    return nested_val

    # Next, any top-level list.
    for val in payload.values():
        if isinstance(val, list):
            return val

    # Finally, one-level nested dict lists.
    for val in payload.values():
        if isinstance(val, dict):
            for nested_val in val.values():
                if isinstance(nested_val, list):
                    return nested_val

    return []


def fetch_all_schools():
    print(f"[→] Fetching schools in {CASELIST}...")
    data = api_get(f"{API_BASE}/caselists/{CASELIST}/schools")
    if not data:
        return []
    schools = _extract_first_list(data, preferred_keys=("schools", "data", "items", "results"))
    print(f"    {len(schools)} schools found")
    return schools


def fetch_teams_in_school(school):
    data = api_get(f"{API_BASE}/caselists/{CASELIST}/schools/{school}/teams")
    if not data:
        return []
    return _extract_first_list(data, preferred_keys=("teams", "data", "items", "results"))


def fetch_recent_rounds(max_pages=RECENT_MAX_PAGES):
    """Fetches recent rounds from the fast caselist endpoint with basic pagination."""
    url = f"{API_BASE}/caselists/{CASELIST}/recent"
    all_rows = []
    seen_ids = set()

    def _extract_rows(payload):
        if isinstance(payload, list):
            return payload
        if isinstance(payload, dict):
            for key in ("recent", "rounds", "data", "items"):
                rows = payload.get(key)
                if isinstance(rows, list):
                    return rows
        return []

    def _add_rows(rows):
        if not rows or not isinstance(rows, list):
            return 0
        added = 0
        for r in rows:
            rid = r.get("round_id")
            if rid in seen_ids:
                continue
            seen_ids.add(rid)
            all_rows.append(r)
            added += 1
        return added

    # First request with no query params is the most reliable form.
    first = []
    try:
        r = session.get(url, timeout=15)
        if r.status_code == 200:
            first = _extract_rows(r.json())
        else:
            print(f"    [recent] first request status {r.status_code}")
    except Exception:
        first = []

    first_added = _add_rows(first)
    if first_added == 0:
        # Fallbacks in case API expects explicit paging args.
        first = _extract_rows(api_get(url, params={"page": 1}) or [])
        if not first:
            first = _extract_rows(api_get(url, params={"offset": 0}) or [])
        first_added = _add_rows(first)

    if first_added == 0:
        return []

    print(f"    [recent] page 1: {first_added} rows")

    # Probe pagination style. Some deployments support page=2, some do not.
    page2 = _extract_rows(api_get(url, params={"page": 2}) or [])
    if page2:
        added = _add_rows(page2)
        print(f"    [recent] page 2: {added} rows")
        for page in range(3, max_pages + 1):
            rows = _extract_rows(api_get(url, params={"page": page}) or [])
            if not rows:
                break
            added = _add_rows(rows)
            print(f"    [recent] page {page}: {added} rows")
            if len(rows) < 50:
                break
    else:
        # Fallback to offset paging if available.
        for offset in range(50, 50 * max_pages, 50):
            rows = _extract_rows(api_get(url, params={"offset": offset}) or [])
            if not rows:
                break
            added = _add_rows(rows)
            print(f"    [recent] offset {offset}: {added} rows")
            if len(rows) < 50:
                break

    return all_rows


def _coalesce_text(*values):
    for value in values:
        if value is None:
            continue
        text = str(value).strip()
        if text:
            return text
    return ""


def _extract_school_team_from_recent_row(row):
    school = _coalesce_text(
        row.get("school_name"),
        row.get("school"),
        row.get("school_slug"),
        row.get("aff_school"),
        row.get("neg_school"),
    )
    team = _coalesce_text(
        row.get("team_name"),
        row.get("team"),
        row.get("team_slug"),
        row.get("aff_team"),
        row.get("neg_team"),
    )

    school_obj = row.get("school")
    if isinstance(school_obj, dict):
        school = school or _coalesce_text(school_obj.get("name"), school_obj.get("slug"))

    team_obj = row.get("team")
    if isinstance(team_obj, dict):
        team = team or _coalesce_text(team_obj.get("name"), team_obj.get("slug"))

    # Last-resort parse from canonical path: caselist/school/team/file.docx
    path = _coalesce_text(row.get("opensource"))
    parts = [p for p in path.split("/") if p]
    if len(parts) >= 4:
        school = school or parts[1]
        team = team or parts[2]

    return school, team


def resolve_recent_by_recent_endpoint(cutoff):
    """Fallback path that groups /recent rows by (school, team)."""
    rows = fetch_recent_rounds()
    if not rows:
        return []

    grouped = defaultdict(list)
    skipped_missing_team = 0

    for row in rows:
        if not _is_recent(row, cutoff):
            continue
        if not _matches_topic(row):
            continue

        school, team = _extract_school_team_from_recent_row(row)
        if not school or not team:
            skipped_missing_team += 1
            continue
        grouped[(school, team)].append(row)

    results = [
        (school, team, rounds)
        for (school, team), rounds in sorted(grouped.items(), key=lambda item: (item[0][0].lower(), item[0][1].lower()))
    ]

    print(
        f"[→] Recent endpoint fallback resolved {len(results)} teams "
        f"from {len(rows)} recent rows."
    )
    if skipped_missing_team:
        print(f"    [recent] skipped {skipped_missing_team} rows with missing school/team metadata")

    return results


def fetch_rounds(school, team):
    cache_key = hashlib.md5(f"{CASELIST}{school}{team}".encode()).hexdigest()
    cache_file = CACHE_DIR / f"rounds_{cache_key}.json"
    if cache_file.exists() and (time.time() - cache_file.stat().st_mtime) < 3600:
        return json.loads(cache_file.read_text())

    # Try two URL patterns
    data = api_get(f"{API_BASE}/caselists/{CASELIST}/schools/{school}/teams/{team}/rounds")
    if data is None:
        data = api_get(f"{API_BASE}/caselists/{CASELIST}/teams/{school}/{team}/rounds")
    if not data:
        return []

    rounds = _extract_first_list(data, preferred_keys=("rounds", "data", "items", "results"))
    cache_file.write_text(json.dumps(rounds))
    time.sleep(0.3)
    return rounds


def resolve_recent_by_school_scan(cutoff):
    """Fallback path when fast /recent endpoint is unavailable."""
    results = []
    schools = fetch_all_schools()
    if not schools:
        print("[!] Could not fetch schools (token expired, rate-limited, or API temporarily unavailable).")
        return []

    total_schools = len(schools)
    scanned_schools = 0
    scanned_teams = 0
    matched_teams = 0

    for school_obj in schools:
        school = school_obj if isinstance(school_obj, str) else school_obj.get("name", "")
        if not school:
            continue
        scanned_schools += 1
        teams = fetch_teams_in_school(school)
        if scanned_schools % SCAN_LOG_EVERY_SCHOOLS == 0:
            print(
                f"    [scan] school {scanned_schools}/{total_schools}: {school} "
                f"({len(teams)} teams)",
                flush=True,
            )

        for team_obj in teams:
            team = team_obj if isinstance(team_obj, str) else (team_obj.get("name") or team_obj.get("team", ""))
            if not team:
                continue
            scanned_teams += 1
            if SCAN_SHOW_TEAM and scanned_teams % 25 == 0:
                print(
                    f"      [team] {school} / {team} (team #{scanned_teams})",
                    flush=True,
                )
            rounds = fetch_rounds(school, team)
            recent = [r for r in rounds if _is_recent(r, cutoff)]
            if recent:
                matched_teams += 1
                results.append((school, team, recent))

        if scanned_schools % 25 == 0 or scanned_schools == total_schools:
            print(
                f"    [progress] schools {scanned_schools}/{total_schools} | "
                f"teams scanned {scanned_teams} | matches {matched_teams}"
            )
        time.sleep(0.2)

    print(
        f"[→] Fallback recent scan complete: {scanned_schools} schools, "
        f"{scanned_teams} teams scanned, {matched_teams} matching teams."
    )
    return results


# ───────────────────────────────────────────────────────────────
#  TARGET RESOLUTION
# ───────────────────────────────────────────────────────────────

# runtime-selected
TARGET_MODE = "teams"  # will be overwritten by prompt


def resolve_targets():
    """Returns list of (school, team, rounds)."""
    results = []

    if TARGET_MODE == "teams":
        total_teams = len(SPECIFIC_TEAMS)
        for idx, (school, team) in enumerate(SPECIFIC_TEAMS, start=1):
            print(f"[→] Team {idx}/{total_teams}: {school} / {team}")
            rounds = fetch_rounds(school, team)
            results.append((school, team, rounds))

    elif TARGET_MODE == "school":
        total_schools = len(SPECIFIC_SCHOOLS)
        scanned_teams = 0
        for school_idx, school in enumerate(SPECIFIC_SCHOOLS, start=1):
            print(f"[→] School {school_idx}/{total_schools}: {school}")
            teams = fetch_teams_in_school(school)
            print(f"    [scan] {len(teams)} teams found")
            for team_obj in teams:
                team = team_obj if isinstance(team_obj, str) else (team_obj.get("name") or team_obj.get("team", ""))
                if not team:
                    continue
                scanned_teams += 1
                rounds = fetch_rounds(school, team)
                results.append((school, team, rounds))
                time.sleep(0.2)
            print(
                f"    [progress] schools {school_idx}/{total_schools} | "
                f"teams scanned {scanned_teams} | teams resolved {len(results)}"
            )

    elif TARGET_MODE == "recent":
        cutoff = datetime.now(timezone.utc).replace(tzinfo=None) - timedelta(days=DAYS_RECENT)
        print(f"[→] Rounds uploaded since {cutoff.strftime('%Y-%m-%d')} ({DAYS_RECENT} days, created_at only)...")
        schools = fetch_all_schools()
        if not schools:
            print("[!] Could not fetch schools (token expired, rate-limited, or API temporarily unavailable).")
            return []
        total_schools = len(schools)
        scanned_schools = 0
        scanned_teams = 0
        matched_teams = 0
        for school_obj in schools:
            school = school_obj if isinstance(school_obj, str) else school_obj.get("name", "")
            if not school:
                continue
            scanned_schools += 1
            teams = fetch_teams_in_school(school)
            if scanned_schools % SCAN_LOG_EVERY_SCHOOLS == 0:
                print(
                    f"    [scan] school {scanned_schools}/{total_schools}: {school} "
                    f"({len(teams)} teams)",
                    flush=True,
                )

            for team_obj in teams:
                team = team_obj if isinstance(team_obj, str) else (team_obj.get("name") or team_obj.get("team", ""))
                if not team:
                    continue
                scanned_teams += 1
                if SCAN_SHOW_TEAM and scanned_teams % 25 == 0:
                    print(
                        f"      [team] {school} / {team} (team #{scanned_teams})",
                        flush=True,
                    )
                rounds = fetch_rounds(school, team)
                recent = [r for r in rounds if _is_recent(r, cutoff)]
                if recent:
                    results.append((school, team, recent))
                    matched_teams += 1

            if scanned_schools % 25 == 0 or scanned_schools == total_schools:
                print(
                    f"    [progress] schools {scanned_schools}/{total_schools} | "
                    f"teams scanned {scanned_teams} | matching teams {matched_teams}"
                )
            time.sleep(0.2)

        print(
            f"[→] Recent scan complete: {scanned_schools} schools, "
            f"{scanned_teams} teams scanned, {matched_teams} matching teams."
        )

    elif TARGET_MODE == "topic":
        if not TOPIC_KEYWORDS:
            print("[!] topic mode requires TOPIC_KEYWORDS to be set!")
            return []
        print(f"[→] Topic scan: {TOPIC_KEYWORDS}")
        schools = fetch_all_schools()
        if not schools:
            print("[!] Could not fetch schools (token expired, rate-limited, or API temporarily unavailable).")
            return []
        total_schools = len(schools)
        scanned_schools = 0
        scanned_teams = 0
        matched_teams = 0
        for school_obj in schools:
            school = school_obj if isinstance(school_obj, str) else school_obj.get("name", "")
            if not school:
                continue
            scanned_schools += 1
            teams = fetch_teams_in_school(school)
            if scanned_schools % SCAN_LOG_EVERY_SCHOOLS == 0:
                print(
                    f"    [scan] school {scanned_schools}/{total_schools}: {school} "
                    f"({len(teams)} teams)",
                    flush=True,
                )

            for team_obj in teams:
                team = team_obj if isinstance(team_obj, str) else (team_obj.get("name") or team_obj.get("team", ""))
                if not team:
                    continue
                scanned_teams += 1
                if SCAN_SHOW_TEAM and scanned_teams % 25 == 0:
                    print(
                        f"      [team] {school} / {team} (team #{scanned_teams})",
                        flush=True,
                    )
                rounds = fetch_rounds(school, team)
                matching = [r for r in rounds if _matches_topic(r)]
                if matching:
                    matched_teams += 1
                    results.append((school, team, matching))

            if scanned_schools % 25 == 0 or scanned_schools == total_schools:
                print(
                    f"    [progress] schools {scanned_schools}/{total_schools} | "
                    f"teams scanned {scanned_teams} | matches {matched_teams}"
                )
            time.sleep(0.2)

        print(
            f"[→] Topic scan complete: {scanned_schools} schools, "
            f"{scanned_teams} teams scanned, {matched_teams} matching teams."
        )

    return results


def _parse_api_timestamp(raw_ts):
    text = (raw_ts or "").strip()
    if not text:
        return None

    # Normalize common API timestamp variants:
    #   2026-03-29 14:22:11
    #   2026-03-29T14:22:11Z
    #   2026-03-29T14:22:11.123Z
    normalized = text.replace("T", " ").replace("Z", "")
    normalized = re.sub(r"\.\d+", "", normalized)
    normalized = re.sub(r"([+-]\d{2}:\d{2})$", "", normalized).strip()
    normalized = normalized[:19]

    try:
        return datetime.strptime(normalized, "%Y-%m-%d %H:%M:%S")
    except Exception:
        return None


def _normalize_tournament_name(name):
    text = str(name or "").strip().lower()
    if not text:
        return ""
    text = re.sub(r"^[0-9\s\-–_:]+", "", text)
    text = re.sub(r"[^a-z0-9]+", " ", text).strip()
    return text


def _is_recent(rnd, cutoff):
    created_dt = _parse_api_timestamp(rnd.get("created_at"))
    return created_dt is not None and created_dt >= cutoff


def _matches_topic(rnd):
    if not TOPIC_KEYWORDS:
        return True
    text = ((rnd.get("report") or "") + " " + (rnd.get("opensource") or "")).lower()
    return any(kw.lower() in text for kw in TOPIC_KEYWORDS)


def dedup_rounds(rounds):
    """One entry per unique file path, topic-filtered."""
    seen = {}
    for r in rounds:
        path = r.get("opensource")
        if path and path not in seen and _matches_topic(r):
            seen[path] = r
    return list(seen.values())


def classify_block_bucket(meta: dict) -> str:
    """
    Assign each round file to a top-level output bucket.
    Priority order is explicit A2 labels first, then side/path-based Pro/Con.
    """
    combined = " ".join([
        str(meta.get("report") or ""),
        str(meta.get("opensource") or ""),
        str(meta.get("tournament") or ""),
    ]).lower()

    # Normalize common formats: "a/2", "a2", "a-2", "a 2"
    compact = re.sub(r"[^a-z0-9]+", " ", combined)

    if re.search(r"\ba\s*2\s*aff\b", compact) or re.search(r"\ba2\s*aff\b", compact):
        return "A2 Aff"
    if re.search(r"\ba\s*2\s*neg\b", compact) or re.search(r"\ba2\s*neg\b", compact):
        return "A2 Neg"

    # Prefer explicit pro/con case labels when available.
    if re.search(r"\bpro\s*case\b", compact) or re.search(r"\baff\s*case\b", compact):
        return "Pro"
    if re.search(r"\bcon\s*case\b", compact) or re.search(r"\bneg\s*case\b", compact):
        return "Con"

    # Path labels are usually more reliable than side metadata for file type.
    path = str(meta.get("opensource") or "").lower()
    if re.search(r"(^|[-_/])pro([-_/]|$)", path):
        return "Pro"
    if re.search(r"(^|[-_/])con([-_/]|$)", path):
        return "Con"

    side = str(meta.get("side") or "").upper()
    if side == "A":
        return "Pro"
    if side == "N":
        return "Con"

    return "Uncategorized"


def classify_content_type(meta: dict) -> str:
    """
    Classify a file by common debate speech/prep terminology.
    """
    combined = " ".join([
        str(meta.get("report") or ""),
        str(meta.get("opensource") or ""),
        str(meta.get("tournament") or ""),
        str(meta.get("round") or ""),
    ]).lower()
    compact = re.sub(r"[^a-z0-9]+", " ", combined)

    # Explicit prep labels first.
    if re.search(r"\brebuttal\s*ev(idence)?\b", compact) or re.search(r"\bevidence\s*(doc|packet|file|dump)\b", compact):
        return "Evidence/Prep Docs"

    # Summary speeches and final focus (common PF terms).
    if re.search(r"\bfinal\s*focus\b", compact) or re.search(r"\bff\b", compact) or re.search(r"\bsummary\b", compact):
        return "Summary/Final Focus"

    # Cross-examination / crossfire style terms.
    if re.search(r"\bcross\s*(ex|examination|fire)\b", compact) or re.search(r"\bcx\b", compact):
        return "Crossfire/CX"

    # Block/rebuttal-heavy terms.
    block_patterns = [
        r"\brebuttal\b",
        r"\bblock\b",
        r"\bfrontline\b",
        r"\bextension\b",
        r"\ba\s*2\b",
        r"\ba2\b",
        r"\b2ac\b",
        r"\b2nc\b",
        r"\b1ar\b",
        r"\b1nr\b",
        r"\b2ar\b",
        r"\b2nr\b",
    ]
    if any(re.search(p, compact) for p in block_patterns):
        return "Blocks/Rebuttals"

    # Constructive/case terms.
    case_patterns = [
        r"\b1ac\b",
        r"\b1nc\b",
        r"\bconstructive\b",
        r"\bcase\b",
        r"\bpro\s*constructive\b",
        r"\bcon\s*constructive\b",
        r"\bpc\b",
        r"\bcc\b",
    ]
    if any(re.search(p, compact) for p in case_patterns):
        return "Cases/Constructives"

    return "General Round Files"


def _add_bucket_heading(doc: Document, label: str):
    h = doc.add_heading(label, level=1)
    if h.runs:
        h.runs[0].font.color.rgb = RGBColor(0x11, 0x56, 0x99)


def _add_tournament_heading(doc: Document, tourn_name: str):
    h = doc.add_heading(tourn_name, level=2)
    if h.runs:
        h.runs[0].font.color.rgb = RGBColor(0x1a, 0x5c, 0xa8)


def _add_content_heading(doc: Document, label: str):
    h = doc.add_heading(label, level=2)
    if h.runs:
        h.runs[0].font.color.rgb = RGBColor(0x2a, 0x6a, 0xa8)


def _add_tournament_subheading(doc: Document, tourn_name: str):
    h = doc.add_heading(tourn_name, level=3)
    if h.runs:
        h.runs[0].font.color.rgb = RGBColor(0x1a, 0x5c, 0xa8)


SPEECH_ORDER = {
    "1AC": 10,
    "1NC": 20,
    "2AC": 30,
    "2NC": 40,
    "1AR": 50,
    "1NR": 60,
    "2AR": 70,
    "2NR": 80,
    "Final Focus": 90,
    "Crossfire/CX": 100,
    "Other": 999,
}


def extract_primary_speech_label(meta: dict) -> str:
    """
    Detect the dominant speech label so files can be ordered consistently.
    """
    combined = " ".join([
        str(meta.get("report") or ""),
        str(meta.get("opensource") or ""),
        str(meta.get("round") or ""),
    ]).lower()
    compact = re.sub(r"[^a-z0-9]+", " ", combined)

    label_patterns = [
        ("1AC", [r"\b1ac\b", r"\bpc\b", r"\bpro\s*constructive\b"]),
        ("1NC", [r"\b1nc\b", r"\bcc\b", r"\bcon\s*constructive\b"]),
        ("2AC", [r"\b2ac\b", r"\bpr\b", r"\bpro\s*rebuttal\b"]),
        ("2NC", [r"\b2nc\b", r"\bcr\b", r"\bcon\s*rebuttal\b"]),
        ("1AR", [r"\b1ar\b"]),
        ("1NR", [r"\b1nr\b"]),
        ("2AR", [r"\b2ar\b", r"\bps\b", r"\bpro\s*summary\b"]),
        ("2NR", [r"\b2nr\b", r"\bcs\b", r"\bcon\s*summary\b"]),
        ("Final Focus", [r"\bfinal\s*focus\b", r"\bff\b", r"\bpf\b", r"\bcf\b"]),
        ("Crossfire/CX", [r"\bcross\s*(ex|examination|fire)\b", r"\bcx\b"]),
    ]

    for label, patterns in label_patterns:
        if any(re.search(p, compact) for p in patterns):
            return label
    return "Other"


def _round_sort_key(round_value: str):
    text = str(round_value or "").strip().lower()
    if text.isdigit():
        return (0, int(text), "")

    elim_order = {
        "doubles": 200,
        "triples": 210,
        "octas": 220,
        "quads": 230,
        "quarters": 240,
        "semis": 250,
        "finals": 260,
        "all": 999,
    }
    if text in elim_order:
        return (1, elim_order[text], "")

    num_match = re.search(r"\d+", text)
    if num_match:
        return (0, int(num_match.group()), text)

    return (2, 999, text)


def _entry_sort_key(meta: dict):
    speech = extract_primary_speech_label(meta)
    rank = SPEECH_ORDER.get(speech, 999)
    round_key = _round_sort_key(str(meta.get("round") or ""))
    fname = Path(str(meta.get("opensource") or "")).name.lower()
    return (rank, round_key, fname)


def _add_file_heading(doc: Document, meta: dict, speech_label: str):
    school = str(meta.get("school") or "")
    team = str(meta.get("team") or "")
    rnd = str(meta.get("round") or "")
    title = f"{school}/{team}  ·  {speech_label}"
    if rnd:
        title += f"  ·  Round {rnd}"
    h = doc.add_heading(title, level=4)
    if h.runs:
        h.runs[0].font.color.rgb = RGBColor(0x35, 0x65, 0x95)


def _demote_imported_heading_style(paragraph_element, paragraph_text=""):
    """
    Preserve heading text but flatten imported heading styles so the compiled
    navigator is driven by our explicit sort hierarchy.
    """
    p_pr = paragraph_element.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        paragraph_element.insert(0, p_pr)

    # Force imported paragraphs to Normal style so source heading styles
    # do not pollute the compiled document's navigation hierarchy.
    p_style = p_pr.find(qn("w:pStyle"))
    if p_style is None:
        p_style = OxmlElement("w:pStyle")
        p_pr.insert(0, p_style)

    old_style_val = p_style.get(qn("w:val"), "")
    was_heading_style = re.fullmatch(r"Heading([1-9])", old_style_val, flags=re.IGNORECASE) is not None
    p_style.set(qn("w:val"), "Normal")

    # Remove explicit outline levels that can still surface in nav trees.
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is not None:
        p_pr.remove(outline)

    # Add deep outline levels for tag/card lines so they appear as indented
    # navigation subsections while keeping body formatting normal.
    if _looks_like_tag_or_card_heading(paragraph_text, was_heading_style):
        outline = OxmlElement("w:outlineLvl")
        outline.set(qn("w:val"), "6")
        p_pr.append(outline)


def _looks_like_tag_or_card_heading(text: str, was_heading_style: bool) -> bool:
    raw = (text or "").strip()
    if not raw:
        return False
    if len(raw) > 180:
        return False

    compact = re.sub(r"[^a-z0-9: ]+", " ", raw.lower()).strip()
    if not compact:
        return False

    # Keep generic packet/case headers out of nav noise.
    if re.search(r"\b(pro|con|aff|neg)\s*case\b", compact):
        return False
    if re.search(r"\baff\s*vs\b", compact):
        return False

    tag_patterns = [
        r"^\d+\s*[\]\).:-]\s*(at|a2|t|nl|nu|turn|ov|fw|link|impact|perm|alt|da|cp|k|solv|case|contention)\s*:",
        r"^(at|a2|t|nl|nu|turn|ov|fw|link|impact|perm|alt|da|cp|k|solv|case|contention)\s*:",
        r"^contention\s*\d+",
    ]
    if any(re.search(p, compact) for p in tag_patterns):
        return True

    # If source used a heading and it looks like a short tag title, keep it in nav.
    if was_heading_style and ":" in raw and len(raw.split()) <= 12:
        return True

    return False


def _infer_side_code(meta: dict) -> str:
    path = str(meta.get("opensource") or "").lower()
    if re.search(r"(^|[-_/])pro([-_/]|$)", path):
        return "A"
    if re.search(r"(^|[-_/])con([-_/]|$)", path):
        return "N"

    side = str(meta.get("side") or "").upper()
    if side in ("A", "N"):
        return side

    compact = re.sub(r"[^a-z0-9]+", " ", " ".join([
        str(meta.get("report") or ""),
        str(meta.get("tournament") or ""),
    ]).lower())
    if re.search(r"\bpro\s*case\b|\baff\s*case\b", compact):
        return "A"
    if re.search(r"\bcon\s*case\b|\bneg\s*case\b", compact):
        return "N"
    return ""


def detect_speech_label_from_text(text: str, side_code: str) -> str:
    """
    Map section-heading text to a canonical speech bucket.
    Explicit labels win; generic labels (constructive/rebuttal/case/summary)
    are mapped by side when possible.
    """
    compact = re.sub(r"[^a-z0-9]+", " ", (text or "").lower()).strip()
    if not compact:
        return ""

    # Explicit speech labels.
    explicit = [
        ("1AC", [r"\b1ac\b", r"\bpc\b", r"\bpro\s*constructive\b"]),
        ("1NC", [r"\b1nc\b", r"\bcc\b", r"\bcon\s*constructive\b"]),
        ("2AC", [r"\b2ac\b", r"\bpr\b", r"\bpro\s*rebuttal\b"]),
        ("2NC", [r"\b2nc\b", r"\bcr\b", r"\bcon\s*rebuttal\b"]),
        ("1AR", [r"\b1ar\b"]),
        ("1NR", [r"\b1nr\b"]),
        ("2AR", [r"\b2ar\b", r"\bps\b", r"\bpro\s*summary\b"]),
        ("2NR", [r"\b2nr\b", r"\bcs\b", r"\bcon\s*summary\b"]),
        ("Final Focus", [r"\bfinal\s*focus\b", r"\bpf\b", r"\bcf\b", r"\bff\b"]),
        ("Crossfire/CX", [r"\bcross\s*(ex|examination|fire)\b", r"\bcx\b"]),
    ]
    for label, patterns in explicit:
        if any(re.search(p, compact) for p in patterns):
            return label

    # Generic headings mapped by side.
    if re.search(r"\b(rebuttal|block|frontline|a\s*2|a2)\b", compact):
        if side_code == "A":
            return "2AC"
        if side_code == "N":
            return "2NC"
        return "Other"

    if re.search(r"\b(constructive|case|pro\s*case|con\s*case)\b", compact):
        if side_code == "A":
            return "1AC"
        if side_code == "N":
            return "1NC"
        return "Other"

    if re.search(r"\bsummary\b", compact):
        if side_code == "A":
            return "2AR"
        if side_code == "N":
            return "2NR"
        return "Other"

    return ""


def _is_speech_boundary_heading(para) -> bool:
    text = (para.text or "").strip()
    if not text:
        return False
    if len(text) > 120:
        return False

    try:
        style_name = (para.style.name or "").lower()
    except Exception:
        style_name = ""

    if style_name.startswith("heading"):
        return True

    # Short title-like lines can also act as section boundaries.
    if len(text.split()) <= 8 and re.fullmatch(r"[A-Za-z0-9\-: /&()]+", text):
        return True
    return False


def split_docx_into_speech_sections(src_bytes: bytes, meta: dict):
    """
    Split a source document into ordered speech sections based on internal headings.
    Returns list of tuples:
      (speech_label, section_title, paragraph_elements)
    """
    try:
        src = Document(io.BytesIO(src_bytes))
    except Exception:
        return []

    side_code = _infer_side_code(meta)
    current_label = extract_primary_speech_label(meta)
    current_title = current_label
    current_paragraphs = []
    sections = []

    for para in src.paragraphs:
        text = (para.text or "").strip()
        detected = ""
        if text and _is_speech_boundary_heading(para):
            detected = detect_speech_label_from_text(text, side_code)

        if detected:
            if detected != current_label:
                if current_paragraphs:
                    sections.append((current_label, current_title, current_paragraphs))
                    current_paragraphs = []
                current_label = detected
                current_title = text[:120]
            elif current_title == current_label:
                # Keep same bucket but upgrade generic title when a better one appears.
                current_title = text[:120]

        new_p = copy.deepcopy(para._element)
        _demote_imported_heading_style(new_p, text)
        current_paragraphs.append(new_p)

    if current_paragraphs:
        sections.append((current_label, current_title, current_paragraphs))

    return sections


def copy_section_into(paragraph_elements, dest_doc: Document, meta: dict,
                      speech_label: str, section_title: str) -> int:
    """
    Inserts a pre-split paragraph section without extra generated wrappers.
    """
    dest_body = dest_doc.element.body
    insert_idx = len(dest_body) - 1
    count = 0
    for p in paragraph_elements:
        new_p = copy.deepcopy(p)
        dest_body.insert(insert_idx, new_p)
        insert_idx += 1
        count += 1

    dest_doc.add_paragraph()
    return count


def _section_entry_sort_key(meta: dict, section_index: int):
    return (_round_sort_key(str(meta.get("round") or "")), section_index,
            Path(str(meta.get("opensource") or "")).name.lower())


def normalize_speech_for_bucket(bucket: str, speech_label: str) -> str:
    """
    Keep speech buckets consistent with side buckets when labels conflict.
    """
    if bucket in ("Con", "A2 Neg"):
        if speech_label == "1AC":
            return "1NC"
        if speech_label == "2AC":
            return "2NC"
    if bucket in ("Pro", "A2 Aff"):
        if speech_label == "1NC":
            return "1AC"
        if speech_label == "2NC":
            return "2AC"
    return speech_label


# ───────────────────────────────────────────────────────────────
#  FILE DOWNLOAD
# ───────────────────────────────────────────────────────────────

def download_file(path: str):
    import tempfile
    import sys
    import os
    
    is_pdf = path.lower().endswith(".pdf")
    
    key = hashlib.md5(path.encode()).hexdigest()
    cached = CACHE_DIR / f"{key}.docx"
    if cached.exists() and cached.stat().st_size > 0:
        return cached.read_bytes()

    print(f"    [↓] {Path(path).name}")
    for attempt in range(3):
        try:
            r = session.get(f"{API_BASE}/download", params={"path": path}, timeout=30)
            if r.status_code == 429:
                wait = 2 ** attempt
                print(f"    [rate limit] {Path(path).name} retrying in {wait}s")
                time.sleep(wait)
                continue
            
            if r.status_code == 200:
                if is_pdf:
                    if not r.content.startswith(b'%PDF'):
                        print(f"    [!] Invalid PDF format for {Path(path).name}")
                        return None
                        
                    try:
                        from pdf2docx import Converter
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
                            tmp_pdf.write(r.content)
                            tmp_pdf_path = tmp_pdf.name
                            
                        tmp_docx_path = tmp_pdf_path + ".docx"
                        
                        # Suppress verbose output from pdf2docx
                        old_stdout = sys.stdout
                        sys.stdout = open(os.devnull, 'w')
                        try:
                            cv = Converter(tmp_pdf_path)
                            cv.convert(tmp_docx_path, start=0, end=None)
                            cv.close()
                        finally:
                            sys.stdout.close()
                            sys.stdout = old_stdout
                        
                        converted_bytes = Path(tmp_docx_path).read_bytes()
                        cached.write_bytes(converted_bytes)
                        
                        Path(tmp_pdf_path).unlink(missing_ok=True)
                        Path(tmp_docx_path).unlink(missing_ok=True)
                        
                        time.sleep(0.6)
                        return converted_bytes
                    except Exception as e:
                        # Restore stdout if exception happened outside the finally block (unlikely)
                        if sys.stdout != old_stdout:
                            sys.stdout.close()
                            sys.stdout = old_stdout
                        print(f"    [!] PDF Conversion failed for {Path(path).name}: {e}")
                        return None
                else:
                    if r.content[:4] == b'PK\x03\x04':
                        cached.write_bytes(r.content)
                        time.sleep(0.6)
                        return r.content
                    else:
                        print(f"    [!] Non-DOCX content returned for {Path(path).name}")
                        return None
            else:
                if attempt == 2:
                    print(f"    [!] Download returned status={r.status_code} for {Path(path).name}")
                time.sleep(2 ** attempt)
        except Exception:
            time.sleep(2 ** attempt)
    print(f"    [!] Failed after 3 attempts: {Path(path).name}")
    return None


# ───────────────────────────────────────────────────────────────
#  FORMAT-PRESERVING DOCX MERGE
# ───────────────────────────────────────────────────────────────

def _add_attr_paragraph(doc, text, hex_color, bold=False, size_pt=10,
                        space_before_pt=0, space_after_pt=3):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before_pt)
    para.paragraph_format.space_after  = Pt(space_after_pt)
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    r = int(hex_color[0:2], 16)
    g = int(hex_color[2:4], 16)
    b = int(hex_color[4:6], 16)
    run.font.color.rgb = RGBColor(r, g, b)
    return para


def _add_rule(doc, color="3366AA"):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(1)
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def copy_docx_into(src_bytes: bytes, dest_doc: Document, meta: dict) -> int:
    """
    Inserts attribution header then copies every paragraph from src_bytes
    into dest_doc using raw XML copy for full format preservation.
    """
    try:
        src = Document(io.BytesIO(src_bytes))
    except Exception as e:
        print(f"    [!] Parse error: {e}")
        return 0

    side  = "AFF" if meta.get("side") == "A" else "NEG"
    tourn = meta.get("tournament", "").lstrip("0123456789-– ").strip()
    rnd   = meta.get("round", "")
    opp   = meta.get("opponent", "")
    judge = meta.get("judge", "")
    fname = Path(meta.get("opensource", "")).name

    # Attribution block
    _add_attr_paragraph(dest_doc, "─" * 72, "2255AA",
                        size_pt=7, space_before_pt=12, space_after_pt=1)
    _add_attr_paragraph(dest_doc,
        f"{meta.get('school','')}  /  {meta.get('team','')}   ·   "
        f"{side}   ·   {tourn}  —  Round {rnd}",
        "1a5fa8", bold=True, size_pt=11, space_before_pt=1, space_after_pt=1)
    if opp:
        _add_attr_paragraph(dest_doc,
            f"vs {opp}   |   Judge: {judge}",
            "777777", size_pt=9, space_after_pt=1)
    report = meta.get("report", "")
    if report:
        _add_attr_paragraph(dest_doc,
            report.replace("\n", "  |  "),
            "999999", size_pt=8, space_after_pt=1)
    _add_attr_paragraph(dest_doc, f"File: {fname}",
                        "AAAAAA", size_pt=7, space_after_pt=3)
    _add_rule(dest_doc)

    # Copy raw XML paragraphs
    dest_body = dest_doc.element.body
    insert_idx = len(dest_body) - 1  # before sectPr
    count = 0
    for para in src.paragraphs:
        new_p = copy.deepcopy(para._element)
        _demote_imported_heading_style(new_p, para.text or "")
        dest_body.insert(insert_idx, new_p)
        insert_idx += 1
        count += 1

    dest_doc.add_paragraph()  # spacing between files
    return count


# ───────────────────────────────────────────────────────────────
#  COVER PAGE
# ───────────────────────────────────────────────────────────────

def build_cover(doc, target_summary, file_count, topic_info):
    h0 = doc.add_heading("OpenCaselist Block Compilation", 0)
    h0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if h0.runs:
        h0.runs[0].font.color.rgb = RGBColor(0x1a, 0x5f, 0xa8)

    doc.add_paragraph()

    for label, value in [
        ("Caselist",      CASELIST),
        ("Mode",          TARGET_MODE),
        ("Targets",       target_summary),
        ("Files",         str(file_count) + " unique round documents"),
        ("Topic filter",  topic_info),
        ("Generated",     datetime.now().strftime("%Y-%m-%d %H:%M")),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lb = p.add_run(f"{label}:  ")
        lb.bold = True
        lb.font.size = Pt(12)
        vl = p.add_run(value)
        vl.font.size = Pt(12)

    doc.add_page_break()


# ───────────────────────────────────────────────────────────────
#  PDF CONVERSION
# ───────────────────────────────────────────────────────────────

def convert_to_pdf(docx_path: Path):
    pdf_path = docx_path.with_suffix(".pdf")

    # Option 1: docx2pdf (Windows + MS Word)
    try:
        from docx2pdf import convert
        print("[→] Converting to PDF via Microsoft Word (docx2pdf)...")
        convert(str(docx_path), str(pdf_path))
        if pdf_path.exists():
            print(f"[✓] PDF saved: {pdf_path.resolve()}")
            return
    except ImportError:
        print("[!] docx2pdf not installed — run: pip install docx2pdf")
    except Exception as e:
        print(f"[!] docx2pdf error: {e}")

    # Option 2: LibreOffice headless
    for cmd in ["soffice", "libreoffice",
                r"C:\Program Files\LibreOffice\program\soffice.exe"]:
        try:
            res = subprocess.run(
                [cmd, "--headless", "--convert-to", "pdf",
                 "--outdir", str(docx_path.parent), str(docx_path)],
                capture_output=True, timeout=120
            )
            if res.returncode == 0 and pdf_path.exists():
                print(f"[✓] PDF saved via LibreOffice: {pdf_path.resolve()}")
                return
        except FileNotFoundError:
            continue
        except Exception:
            continue

    print("\n" + "=" * 55)
    print("  DOCX saved but PDF conversion unavailable.")
    print("  To get a PDF, either:")
    print("    1.  pip install docx2pdf  (needs MS Word)")
    print("    2.  Open the .docx in Word → Save As → PDF")
    print(f"\n  DOCX is at: {docx_path.resolve()}")
    print("=" * 55 + "\n")


# ───────────────────────────────────────────────────────────────
#  MAIN
# ───────────────────────────────────────────────────────────────

def main():
    global TARGET_MODE, SPECIFIC_TEAMS, SPECIFIC_SCHOOLS, DAYS_RECENT, TOPIC_KEYWORDS

    print(f"\n{'='*60}")
    print("  OpenCaselist Scraper v2 (Interactive)")
    print(f"  caselist={CASELIST}")
    print(f"{'='*60}")

    if DOWNLOAD_ONLY:
        if not LAST_METAS_FILE.exists():
            print(f"[!] No saved manifest found at: {LAST_METAS_FILE.resolve()}")
            print("[!] Run a normal scan once to create it, then use DOWNLOAD_ONLY=1.")
            return

        try:
            all_metas = json.loads(LAST_METAS_FILE.read_text())
        except Exception as e:
            print(f"[!] Failed to read manifest: {e}")
            return

        if DOWNLOAD_FAILED_ONLY:
            if not LAST_FAILED_FILE.exists():
                print("[!] No previous failed-file list found.")
                return
            try:
                failed_paths = set(json.loads(LAST_FAILED_FILE.read_text()))
            except Exception as e:
                print(f"[!] Failed to read failed-file list: {e}")
                return
            all_metas = [m for m in all_metas if m.get("opensource") in failed_paths]

        print(f"[→] Download-only mode using saved manifest: {len(all_metas)} files")
        if DOWNLOAD_FAILED_ONLY:
            print("[→] Retrying only previously failed files")

        downloaded = []
        failed = []
        total_to_download = len(all_metas)
        for idx, meta in enumerate(all_metas, start=1):
            data = download_file(meta.get("opensource", ""))
            if data:
                downloaded.append((meta, data))
            else:
                if meta.get("opensource"):
                    failed.append(meta["opensource"])
            if idx % DOWNLOAD_LOG_EVERY == 0 or idx == total_to_download:
                print(
                    f"    [download] {idx}/{total_to_download} | "
                    f"ok {len(downloaded)} | failed {len(failed)}",
                    flush=True,
                )

        LAST_FAILED_FILE.write_text(json.dumps(sorted(set(failed)), indent=2))

        print(f"\n[✓] Downloaded: {len(downloaded)}")
        print(f"[!] Failed: {len(set(failed))}")
        print(f"[→] Failed list saved: {LAST_FAILED_FILE.resolve()}")
        return

    # 0) Prompt for mode and config
    TARGET_MODE, updates = prompt_for_target_mode()
    for k, v in updates.items():
        globals()[k] = v

    # Optional: apply topic filter on top of non-topic modes.
    # In topic mode, keywords were already collected in prompt_for_target_mode().
    if TARGET_MODE != "topic":
        extra_filter = prompt_optional_topic_filter()
        if extra_filter is not None:
            TOPIC_KEYWORDS = extra_filter

    print(f"\n[→] Running with mode={TARGET_MODE}")
    if TOPIC_KEYWORDS:
        print(f"[→] Topic filter: {TOPIC_KEYWORDS}")

    # 1) Resolve targets
    team_data = resolve_targets()
    if not team_data:
        print("[!] No targets resolved. Try broader keywords, a larger recent window, or teams/school mode.")
        return
    print(f"\n[✓] {len(team_data)} teams resolved\n")

    # 2) Collect unique files per team
    all_metas = []
    for (school, team, rounds) in team_data:
        unique = dedup_rounds(rounds)
        print(f"  {school}/{team}: {len(unique)} unique files")
        for rnd in unique:
            if "opensource" not in rnd or not rnd["opensource"]:
                continue
            all_metas.append({
                "school":     school,
                "team":       team,
                "tournament": rnd.get("tournament", ""),
                "round":      rnd.get("round", ""),
                "side":       rnd.get("side", ""),
                "opponent":   rnd.get("opponent", ""),
                "judge":      rnd.get("judge", ""),
                "report":     rnd.get("report", ""),
                "opensource": rnd["opensource"],
                "created_at": rnd.get("created_at", ""),
            })

    print(f"\n[→] Downloading {len(all_metas)} files...\n")
    LAST_METAS_FILE.write_text(json.dumps(all_metas, indent=2))
    print(f"[→] Saved manifest: {LAST_METAS_FILE.resolve()}")
    downloaded = []
    failed = []
    total_to_download = len(all_metas)
    for idx, meta in enumerate(all_metas, start=1):
        data = download_file(meta["opensource"])
        if data:
            downloaded.append((meta, data))
        else:
            failed.append(meta["opensource"])
        if idx % DOWNLOAD_LOG_EVERY == 0 or idx == total_to_download:
            print(
                f"    [download] {idx}/{total_to_download} | "
                f"ok {len(downloaded)} | failed {len(failed)}",
                flush=True,
            )

    LAST_FAILED_FILE.write_text(json.dumps(sorted(set(failed)), indent=2))
    print(f"[→] Failed list saved: {LAST_FAILED_FILE.resolve()}")

    print(f"\n[✓] {len(downloaded)} files ready\n")
    if not downloaded:
        print("[!] Nothing to compile.")
        return

    # 3) Build output DOCX
    print("[→] Building output document (original formatting preserved)...")
    out_doc = Document()
    for section in out_doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.85)
        section.right_margin  = Inches(0.85)

    # Cover
    target_summary = (
        ", ".join(f"{s}/{t}" for s, t, _ in team_data)
        if len(team_data) <= 5 else f"{len(team_data)} teams"
    )
    topic_info = (
        " | ".join(TOPIC_KEYWORDS) if TOPIC_KEYWORDS
        else "none (all rounds included)"
    )
    build_cover(out_doc, target_summary, len(downloaded), topic_info)

    # Group by bucket (A2/side), then speech bucket, then tournament.
    by_bucket_speech_tourn = defaultdict(lambda: defaultdict(lambda: defaultdict(list)))
    for (meta, data) in downloaded:
        bucket = classify_block_bucket(meta)
        tourn = meta["tournament"].lstrip("0123456789-– ").strip() or "Unknown"
        sections = split_docx_into_speech_sections(data, meta)
        if not sections:
            sections = [(extract_primary_speech_label(meta), "Full File", [])]

        for idx, (speech_label, section_title, paragraph_elements) in enumerate(sections):
            speech_label = normalize_speech_for_bucket(bucket, speech_label)
            by_bucket_speech_tourn[bucket][speech_label][tourn].append(
                (meta, idx, section_title, paragraph_elements)
            )

    bucket_order = ["A2 Aff", "A2 Neg", "Pro", "Con", "Uncategorized"]
    speech_order = [
        "1AC",
        "1NC",
        "2AC",
        "2NC",
        "1AR",
        "1NR",
        "2AR",
        "2NR",
        "Final Focus",
        "Crossfire/CX",
        "Other",
    ]
    present_buckets = [b for b in bucket_order if b in by_bucket_speech_tourn]

    print("[→] Section counts:")
    for bucket in present_buckets:
        bucket_total = 0
        for tmap in by_bucket_speech_tourn[bucket].values():
            bucket_total += sum(len(v) for v in tmap.values())
        print(f"  - {bucket}: {bucket_total} sections")

    for bucket in present_buckets:
        _add_bucket_heading(out_doc, bucket)
        speech_map = by_bucket_speech_tourn[bucket]
        present_speeches = [s for s in speech_order if s in speech_map]

        for speech_label in present_speeches:
            _add_content_heading(out_doc, speech_label)
            tourn_map = speech_map[speech_label]

            for tourn_name in sorted(tourn_map.keys(), key=lambda x: str(x).lower()):
                entries = sorted(
                    tourn_map[tourn_name],
                    key=lambda row: _section_entry_sort_key(row[0], row[1])
                )
                _add_tournament_subheading(out_doc, tourn_name)

                for (meta, sec_idx, section_title, paragraph_elements) in entries:
                    n = copy_section_into(
                        paragraph_elements, out_doc, meta, speech_label, section_title
                    )
                    print(
                        f"  ✓  [{bucket} | {speech_label}] "
                        f"{Path(meta['opensource']).name}  "
                        f"({n} paragraphs, section #{sec_idx + 1}: {section_title})"
                    )

        out_doc.add_page_break()

    # 4) Save
    docx_path = OUTPUT_DIR / f"{OUTPUT_NAME}.docx"
    out_doc.save(str(docx_path))
    print(f"\n[✓] DOCX saved: {docx_path.resolve()}")

    # 5) PDF
    convert_to_pdf(docx_path)

    print(f"\n{'='*60}")
    print(f"  Done!  Folder: {OUTPUT_DIR.resolve()}")
    print(f"{'='*60}\n")


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n[!] Cancelled by user (Ctrl+C). Exiting cleanly.")
