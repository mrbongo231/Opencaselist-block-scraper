#!/usr/bin/env python3
"""
Promote AT/contention block tags to Heading paragraphs for Google Docs navigation.

Usage:
    python3 google_docs_nav_at_headings.py input.docx
    python3 google_docs_nav_at_headings.py input.docx -o output.docx
    python3 google_docs_nav_at_headings.py input.docx --heading-level 6 --in-place
    python3 google_docs_nav_at_headings.py input.docx --include-contentions --remove-file-headers --at-contention-only
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


AT_PATTERNS = [
    re.compile(r"^\s*AT\s*[:\-]", flags=re.IGNORECASE),
    re.compile(r"^\s*A\s*[/-]?\s*2\s*[:\-]", flags=re.IGNORECASE),
    re.compile(r"^\s*\d+\s*[\].)\-:]\s*(AT|A\s*[/-]?\s*2)\s*[:\-]", flags=re.IGNORECASE),
]

CONTENTION_PATTERNS = [
    re.compile(r"^\s*contention\s*\d*\s*[:\-]", flags=re.IGNORECASE),
    re.compile(r"^\s*c\s*\d+\s*[:\-]", flags=re.IGNORECASE),
]

GENERATED_FILE_HEADER_PATTERNS = [
    re.compile(r"^\s*\d+\.\s+.+\s/\s.+\s\|\s.+\.docx\s*$", flags=re.IGNORECASE),
    re.compile(r"^\s*\d+\.\s+.+\s/\s.+\s\|\s*$", flags=re.IGNORECASE),
]


def looks_like_at_heading(text: str) -> bool:
    raw = (text or "").strip()
    if not raw:
        return False
    if len(raw) > 180:
        return False
    return any(pattern.search(raw) for pattern in AT_PATTERNS)


def looks_like_contention_heading(text: str) -> bool:
    raw = (text or "").strip()
    if not raw:
        return False
    if len(raw) > 180:
        return False
    return any(pattern.search(raw) for pattern in CONTENTION_PATTERNS)


def looks_like_generated_file_header(text: str) -> bool:
    raw = (text or "").strip()
    if not raw:
        return False
    return any(pattern.search(raw) for pattern in GENERATED_FILE_HEADER_PATTERNS)


def ensure_heading_style(doc: Document, level: int) -> str:
    style_name = f"Heading {level}"
    _ = doc.styles[style_name]
    return style_name


def promote_target_lines(doc: Document, heading_style: str, include_contentions: bool) -> int:
    changed = 0
    for para in doc.paragraphs:
        text = para.text or ""
        is_target = looks_like_at_heading(text)
        if include_contentions:
            is_target = is_target or looks_like_contention_heading(text)

        if not is_target:
            continue

        try:
            current_style_name = para.style.name or ""
        except Exception:
            current_style_name = ""

        if current_style_name.lower().startswith("heading"):
            continue

        para.style = heading_style
        changed += 1

    return changed


def remove_generated_file_headers(doc: Document) -> int:
    removed = 0
    for para in list(doc.paragraphs):
        text = para.text or ""
        if not looks_like_generated_file_header(text):
            continue

        p = para._element
        parent = p.getparent()
        if parent is not None:
            parent.remove(p)
            removed += 1

    return removed


def demote_non_target_headings(doc: Document, heading_style: str, include_contentions: bool) -> int:
    demoted = 0
    for para in doc.paragraphs:
        text = para.text or ""

        try:
            current_style_name = para.style.name or ""
        except Exception:
            current_style_name = ""

        if not current_style_name.lower().startswith("heading"):
            continue

        is_target = looks_like_at_heading(text)
        if include_contentions:
            is_target = is_target or looks_like_contention_heading(text)

        if is_target:
            para.style = heading_style
            continue

        para.style = "Normal"
        demoted += 1

    return demoted


def strip_non_target_outline_levels(doc: Document, include_contentions: bool) -> int:
    """
    Remove explicit outline levels from non-target paragraphs.

    This keeps paragraph content untouched while ensuring only AT/contention
    tag lines remain visible in document navigation panes that honor
    w:outlineLvl metadata.
    """
    changed = 0
    for para in doc.paragraphs:
        text = para.text or ""

        is_target = looks_like_at_heading(text)
        if include_contentions:
            is_target = is_target or looks_like_contention_heading(text)

        p = para._p
        ppr = p.find(qn("w:pPr"))
        if ppr is None:
            continue

        outline = ppr.find(qn("w:outlineLvl"))
        if outline is None:
            continue

        if not is_target:
            ppr.remove(outline)
            changed += 1

    return changed


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Promote AT/A2 tag lines in a DOCX to headings for Google Docs nav."
    )
    parser.add_argument("input", type=Path, help="Path to input .docx file")
    parser.add_argument(
        "-o",
        "--output",
        type=Path,
        default=None,
        help="Output .docx path (default: *_with_at_nav.docx)",
    )
    parser.add_argument(
        "--heading-level",
        type=int,
        default=6,
        choices=[1, 2, 3, 4, 5, 6],
        help="Heading level to apply to AT lines (default: 6)",
    )
    parser.add_argument(
        "--in-place",
        action="store_true",
        help="Overwrite input file instead of creating a new output file",
    )
    parser.add_argument(
        "--include-contentions",
        action="store_true",
        help="Also promote Contention-style lines as headings",
    )
    parser.add_argument(
        "--remove-file-headers",
        action="store_true",
        help="Remove generated numbered file-label lines like '123. School / Team | File.docx'",
    )
    parser.add_argument(
        "--at-contention-only",
        action="store_true",
        help="Demote non AT/contention heading styles to Normal so outline is cleaner",
    )
    parser.add_argument(
        "--strip-non-target-outline",
        action="store_true",
        help="Remove non AT/contention w:outlineLvl metadata while preserving paragraph text",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()

    input_path = args.input
    if not input_path.exists() or input_path.suffix.lower() != ".docx":
        raise SystemExit("Input must be an existing .docx file")

    if args.in_place:
        output_path = input_path
    elif args.output is not None:
        output_path = args.output
    else:
        output_path = input_path.with_name(f"{input_path.stem}_with_at_nav{input_path.suffix}")

    doc = Document(str(input_path))
    heading_style = ensure_heading_style(doc, args.heading_level)

    removed = 0
    if args.remove_file_headers:
        removed = remove_generated_file_headers(doc)

    changed = promote_target_lines(doc, heading_style, args.include_contentions)

    demoted = 0
    if args.at_contention_only:
        demoted = demote_non_target_headings(doc, heading_style, args.include_contentions)

    outline_stripped = 0
    if args.strip_non_target_outline:
        outline_stripped = strip_non_target_outline_levels(doc, args.include_contentions)

    doc.save(str(output_path))

    print(f"Input:    {input_path}")
    print(f"Output:   {output_path}")
    print(f"Changed:  {changed} promoted AT/A2{'/Contention' if args.include_contentions else ''} headings")
    if args.remove_file_headers:
        print(f"Removed:  {removed} generated file header lines")
    if args.at_contention_only:
        print(f"Demoted:  {demoted} non-target headings")
    if args.strip_non_target_outline:
        print(f"Outline:  {outline_stripped} non-target outline levels removed")
    print(f"Style:    {heading_style}")


if __name__ == "__main__":
    main()
